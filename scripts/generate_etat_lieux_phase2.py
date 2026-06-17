"""
Ouvre reports/Phase1_etat_des_lieux.docx, ajoute l'état des lieux Phase 2
et sauvegarde sous reports/Phases1_2_etat_des_lieux.docx.
"""

import datetime
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
INPUT  = ROOT / "reports" / "Sara" / "Phase1_etat_des_lieux.docx"
OUTPUT = ROOT / "reports" / "Sara" / "Phases1_2_etat_des_lieux.docx"

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def add_info_box(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"ℹ️  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_warning_box(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)


# ── Ouvrir le document existant ───────────────────────────────────────────────
doc = Document(INPUT)

doc.add_page_break()

# ── En-tête Phase 2 ───────────────────────────────────────────────────────────
h = doc.add_heading("État des lieux — Phase 2 : MLflow Tracking & Model Registry", level=2)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph(
    f"Date : {datetime.date.today().strftime('%d %B %Y')}     "
    "Deadline Phase 2 : 1er Juillet 2026"
)

add_warning_box(doc,
    "Ni Airflow ni Cron ne sont utilisés en Phase 2. "
    "Le scheduling automatique des retrains est prévu en Phase 3."
)

doc.add_paragraph()

# ── Tableau état des lieux Phase 2 ───────────────────────────────────────────
doc.add_heading("Composants Phase 2", level=3)

etat_rows = [
    ("1",  "Serveur MLflow (Docker)",
     "✅ OK",
     "docker/mlflow/Dockerfile\ndocker/docker-compose.dev.yml",
     "Container port 5001 — mlflow==3.11.1\n--serve-artifacts activé (upload HTTP)"),
    ("2",  "Tracking training.py",
     "✅ OK",
     "src/train/training.py",
     "10 params + 11 métriques + 3 artifacts\n+ 2 tags loggés à chaque run"),
    ("3",  "Artifacts MLflow",
     "✅ OK",
     "models/confusion_matrix.png\nmodels/classification_report.txt\nmodels/label_mapping.json",
     "Uploadés via HTTP → visibles dans MLflow UI\nsous l'onglet Artifacts du run"),
    ("4",  "Model Registry (aliases)",
     "✅ OK",
     "MLFLOW_MODEL_NAME\n= blood-cell-densenet121",
     "Aliases @production / @challenger\nMLflow 3.x — stages dépréciés"),
    ("5",  "Garde-fous de promotion",
     "✅ OK",
     "training.py\n_register_and_promote()\nRECALL_TOLERANCE = 0.02",
     "macro_f1 >= prod\nrecall_erythroblast et recall_ig\nnon régressifs (tolérance ±2%)"),
    ("6",  "predict_model câblé MLflow",
     "✅ OK",
     "src/models/predict_model.py",
     "Charge @production depuis Registry\nFallback .pth si MLflow indisponible"),
    ("7",  "Logging Supabase prédictions",
     "✅ OK",
     "predict_model\n._log_to_supabase()",
     "inference_ms + predicted_class\nINSERT INTO predictions"),
    ("8",  "MLproject",
     "✅ OK",
     "MLproject\nconda.yaml",
     "mlflow run . -e train --env-manager=local\n-P data_dir=... -P epochs_head=..."),
    ("9",  "Tables SQL MLflow",
     "✅ AUTO",
     "mlruns/mlflow.db (SQLite)\nalembic_version",
     "Créées automatiquement par Alembic\nau démarrage du container Docker"),
    ("10", "Demo garde-fou",
     "✅ OK",
     "scripts/demo_garde_fou.py",
     "Simule un challenger dégradé\nMontre @production protégé dans MLflow UI"),
    ("11", "API /predict câblée MLflow",
     "⏳ Phase 3",
     "src/serving/api.py",
     "ATTENTION : distinct du #6.\napi.py a sa propre logique d'inférence\n"
     "(pas d'import de predict_model.py).\n"
     "Charge le .pth local — pas de Registry.\nRefactoring prévu Phase 3."),
    ("12", "Scheduling retrain",
     "⏳ Phase 3",
     "Airflow / Cron\n(à définir)",
     "Aucun scheduler en Phase 2\nDécision à valider avec le mentor"),
]

STATUS_COLORS = {
    "✅ OK":      ("D4EDDA", RGBColor(0x00, 0x60, 0x00)),
    "✅ AUTO":    ("FFF3CD", RGBColor(0x85, 0x60, 0x00)),
    "⏳ Phase 3": ("F8D7DA", RGBColor(0x80, 0x00, 0x00)),
}

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
for cell, hdr in zip(table.rows[0].cells, ["#", "Livrable", "Statut", "Fichier(s)", "Détail"]):
    set_cell_bg(cell, "1F497D")
    set_cell_border(cell)
    r = cell.paragraphs[0].add_run(hdr)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)

for i, (num, livrable, statut, fichier, detail) in enumerate(etat_rows):
    row = table.add_row().cells
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    status_bg, status_color = STATUS_COLORS.get(statut, (bg, RGBColor(0, 0, 0)))
    for j, (cell, content) in enumerate(zip(row, [num, livrable, statut, fichier, detail])):
        cell_bg = status_bg if j == 2 else bg
        set_cell_bg(cell, cell_bg)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(content)
        r.font.size = Pt(8.5)
        if j == 2:
            r.bold = True
            r.font.color.rgb = status_color

doc.add_paragraph()
add_info_box(doc,
    "Items #6 et #11 — distinction importante :\n"
    "#6 predict_model.py (CLI) est câblé au Registry MLflow (@production) — FAIT.\n"
    "#11 api.py /predict est indépendant : il a sa propre logique d'inférence dupliquée "
    "et ne fait aucun import de predict_model.py. Il charge encore le .pth local "
    "sans passer par MLflow Registry — prévu Phase 3."
)
doc.add_paragraph()

# ── Tables SQL MLflow ─────────────────────────────────────────────────────────
doc.add_heading("Tables SQL MLflow (auto-créées par Alembic)", level=3)
add_info_box(doc,
    "Contrairement à Supabase (init_db.py manuel), les tables MLflow sont créées "
    "automatiquement par le framework au premier démarrage du container."
)

sql_rows = [
    ("experiments",              "Expérience bloodcells-densenet121"),
    ("runs",                     "Chaque run training.py (run_id, status, timestamps)"),
    ("params",                   "batch_size, lr_head, n_train... (mlflow.log_param)"),
    ("metrics / latest_metrics", "macro_f1, recall_ig, recall_erythroblast..."),
    ("tags",                     "git_commit, run_type"),
    ("registered_models",        "blood-cell-densenet121"),
    ("model_versions",           "v1 → vN (une version par run)"),
    ("registered_model_aliases", "@production, @challenger"),
    ("logged_models",            "Modèles via mlflow.pytorch.log_model()"),
    ("alembic_version",          "Version du schéma — migrations automatiques"),
]

table2 = doc.add_table(rows=1, cols=2)
table2.style = "Table Grid"
for cell, hdr in zip(table2.rows[0].cells, ["Table SQLite", "Rôle dans le projet"]):
    set_cell_bg(cell, "2E75B6")
    set_cell_border(cell)
    r = cell.paragraphs[0].add_run(hdr)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)

for i, (tname, role) in enumerate(sql_rows):
    row = table2.add_row().cells
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    for j, (cell, content) in enumerate(zip(row, [tname, role])):
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(content)
        r.font.size = Pt(9)
        if j == 0:
            r.font.name = "Courier New"

doc.add_paragraph()

# ── Garde-fous ────────────────────────────────────────────────────────────────
doc.add_heading("Logique des garde-fous de promotion", level=3)
add_code_block(doc,
    "Nouveau run  ->  @challenger (toujours)\n"
    "       |\n"
    "       v\n"
    "  macro_f1            >= prod_f1               ?\n"
    "  recall_erythroblast >= prod_ery - 0.02        ?   <- tolérance 2%\n"
    "  recall_ig           >= prod_ig  - 0.02        ?   <- tolérance 2%\n"
    "       |\n"
    "  OUI (les 3)              NON (un seul suffit)\n"
    "  @production           reste @challenger + raison affichée"
)

doc.add_paragraph()
add_info_box(doc,
    "Démonstration : python scripts/demo_garde_fou.py\n"
    "Affiche les valeurs réelles de @production et simule un challenger dégradé."
)

# ── Métriques loggées ─────────────────────────────────────────────────────────
doc.add_heading("Métriques et paramètres loggés (chaque run)", level=3)

doc.add_paragraph("Paramètres (10) :").runs[0].bold = True
add_code_block(doc,
    "batch_size  lr_head  lr_full  weight_decay  epochs_head  epochs_full\n"
    "n_train  n_val  n_test  optimizer  part"
)
doc.add_paragraph("Métriques (11 + courbes epoch) :").runs[0].bold = True
add_code_block(doc,
    "best_val_acc  test_acc  macro_f1  weighted_f1  precision_macro  recall_macro\n"
    "recall_erythroblast  recall_ig   <- classes critiques\n"
    "train_time_s  n_params\n"
    "train_loss_epN  val_loss_epN  train_acc_epN  val_acc_epN"
)
doc.add_paragraph("Tags (2) :").runs[0].bold = True
add_code_block(doc, "git_commit   run_type (base | retrain)")

# ── Sauvegarder ───────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Document généré : {OUTPUT}")
