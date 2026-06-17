"""Génère le guide de test Phase 2 (MLflow tracking) en Word dans reports/."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = Path(__file__).parents[1] / "reports" / "Sara" / "Guide_Test_Phase2.docx"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def add_info_box(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"ℹ️  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_warning_box(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)


def add_success_box(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"✅  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)


doc = Document()

# ── Style global ──────────────────────────────────────────────────────────────
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── Titre ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("Projet MLOps — Classification de Cellules Sanguines", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_heading("Guide de Test — Phase 2 : MLflow Tracking (Windows & Mac)", level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    f"Date : {datetime.date.today().strftime('%d %B %Y')}     Deadline Phase 2 : 1er Juillet 2026"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── Contexte Phase 2 ──────────────────────────────────────────────────────────
doc.add_heading("Objectifs Phase 2", level=2)
p = doc.add_paragraph(
    "La Phase 2 ajoute le tracking MLflow complet au pipeline :\n"
    "  • Logging de métriques et paramètres dans chaque run d'entraînement\n"
    "  • Artifacts : matrice de confusion, rapport de classification, mapping des classes\n"
    "  • Tags : commit git, type de run\n"
    "  • Model Registry avec aliases : @production (meilleur modèle) / @challenger (candidat)\n"
    "  • Garde-fous de promotion : F1 macro + recall des classes critiques (erythroblast, ig)\n"
    "  • Logging automatique des prédictions dans Supabase (inference_ms inclus)\n"
    "  • MLproject pour reproductibilité des runs"
)
p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ── Prérequis Phase 2 ─────────────────────────────────────────────────────────
doc.add_heading("Prérequis Phase 2", level=2)

add_info_box(doc, "Les prérequis de la Phase 1 s'appliquent toujours "
             "(environnement virtuel activé, .env configuré, tables Supabase existantes).")

doc.add_paragraph("Prérequis 1 — Docker Desktop installé et démarré").runs[0].bold = True
p = doc.add_paragraph(
    "Télécharger Docker Desktop depuis https://www.docker.com/products/docker-desktop/\n"
    "Vérifier que Docker est actif (icône dans la barre des tâches, statut 'Running')."
)
p.runs[0].font.size = Pt(10)

doc.add_heading("Vérifier Docker (Terminal)", level=3)
add_code_block(doc, "docker --version\n# Attendu : Docker version 24.x.x ou supérieur")

doc.add_paragraph("Prérequis 2 — Port 5001 disponible").runs[0].bold = True
add_code_block(doc,
    "# Mac\nnc -z localhost 5001 && echo 'PORT OCCUPE' || echo 'PORT LIBRE'\n\n"
    "# Windows (PowerShell)\nTest-NetConnection -ComputerName localhost -Port 5001"
)
add_warning_box(doc, "Si le port 5001 est occupé, arrêter l'application qui l'utilise "
                "ou modifier le mapping dans docker/docker-compose.dev.yml.")

doc.add_paragraph("Prérequis 3 — .env avec MLFLOW_TRACKING_URI").runs[0].bold = True
add_code_block(doc, "MLFLOW_TRACKING_URI=http://localhost:5001")
add_warning_box(doc, "Attention : la valeur doit être exactement 'http://localhost:5001' "
                "sans répétition du nom de variable (erreur connue lors de la copie).")

# ── TEST 1 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 1 — Démarrage du serveur MLflow (Docker)", level=2)
add_info_box(doc, "Le serveur MLflow tourne dans un container Docker isolé sur le port 5001. "
             "Les données sont persistées dans mlruns/mlflow.db (SQLite).")

doc.add_heading("Terminal Mac & Windows (PowerShell)", level=3)
add_code_block(doc,
    "# Se placer dans le projet\ncd ~/fev26_bmle_blood_cells          # Mac\n"
    "cd \"$env:USERPROFILE\\fev26_bmle_blood_cells\"  # Windows\n\n"
    "# Construire et démarrer le container MLflow\n"
    "docker-compose -f docker/docker-compose.dev.yml up -d --build mlflow"
)

doc.add_paragraph("Résultat attendu (après ~30s) :").runs[0].bold = True
add_code_block(doc,
    "[+] Building ...\n"
    " => [mlflow] pip install mlflow==3.11.1 psycopg2-binary\n"
    " => Successfully built ...\n"
    "[+] Running 1/1\n"
    " ✔ Container blood_cell_mlflow  Started"
)

doc.add_paragraph("Vérifier que le serveur répond :").runs[0].bold = True
add_code_block(doc,
    "# Mac (curl)\ncurl http://localhost:5001/health\n\n"
    "# Windows (PowerShell)\nInvoke-RestMethod -Uri \"http://localhost:5001/health\" -Method GET"
)
add_code_block(doc, '{"status": "ok"}')

doc.add_paragraph("Interface graphique MLflow :").runs[0].bold = True
p = doc.add_paragraph("Ouvrir dans un navigateur : ")
r = p.add_run("http://localhost:5001")
r.font.name = "Courier New"
r.bold = True
p.add_run(" → page d'accueil MLflow avec la liste des expériences.")

doc.add_paragraph("Arrêter le serveur (optionnel) :").runs[0].bold = True
add_code_block(doc, "docker-compose -f docker/docker-compose.dev.yml stop mlflow")

doc.add_heading("Dépannage Test 1 — Erreurs courantes", level=3)

doc.add_paragraph("Erreur : container name already in use").runs[0].bold = True
add_code_block(doc,
    "Error response from daemon: Conflict. The container name \"/blood_cell_mlflow\"\n"
    "is already in use by container \"xxxxxxxx\"."
)
doc.add_paragraph("Solution : supprimer le container existant puis relancer :")
add_code_block(doc,
    "docker compose -f docker/docker-compose.dev.yml down mlflow\n"
    "docker compose -f docker/docker-compose.dev.yml up -d --build mlflow"
)
add_info_box(doc, "Alternative si tu veux juste redémarrer sans rebuild : "
             "docker compose -f docker/docker-compose.dev.yml restart mlflow")

doc.add_paragraph("Erreur : port 5001 already allocated").runs[0].bold = True
add_code_block(doc,
    "Error response from daemon: driver failed programming external connectivity:\n"
    "Bind for 0.0.0.0:5001 failed: port is already allocated."
)
doc.add_paragraph("Solution : trouver et arrêter le processus qui utilise le port :")
add_code_block(doc,
    "# Windows\nGet-Process -Id (Get-NetTCPConnection -LocalPort 5001).OwningProcess\n"
    "Stop-Process -Id <PID> -Force\n\n"
    "# Mac\nlsof -i :5001\nkill -9 <PID>"
)

# ── TEST 2 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 2 — Entraînement avec tracking MLflow", level=2)
add_warning_box(doc, "Le serveur MLflow (Test 1) doit être démarré avant ce test.")
add_info_box(doc, "Ce test valide le logging MLflow complet : paramètres, métriques epoch par epoch, "
             "artifacts et enregistrement du modèle dans le Registry avec alias @production.")

doc.add_heading("VSCode (Windows & Mac)", level=3)
for step in [
    'Run & Debug → sélectionner "Test training (Source_100 - 1 epoch)"',
    "Cliquer sur ▶️ — durée ~2-4 minutes",
]:
    doc.add_paragraph(step, style="List Number")

doc.add_heading("Terminal Mac", level=3)
add_code_block(doc,
    "python -m src.train.training \\\n"
    "  --data-dir data/Source_100 --output-dir models \\\n"
    "  --epochs-head 1 --epochs-full 1 --batch-size 8 \\\n"
    "  --run-type base --part 0"
)
doc.add_heading("Terminal Windows (PowerShell)", level=3)
add_code_block(doc,
    "python -m src.train.training `\n"
    "  --data-dir data/Source_100 --output-dir models `\n"
    "  --epochs-head 1 --epochs-full 1 --batch-size 8 `\n"
    "  --run-type base --part 0"
)

doc.add_paragraph("Résultat attendu (console) :").runs[0].bold = True
add_code_block(doc,
    "Device  : cpu (ou mps sur Mac M1/M2)\n"
    "Images  : 100 dans 8 classes\n"
    "Split   : train=69  val=16  test=15\n\n"
    "Phase 1 — backbone gelé (1 epoch)\n"
    "  Ep 01  train=x.xxx  val=x.xxx  (Xs)\n\n"
    "Phase 2 — fine-tuning (1 epoch, patience=3)\n"
    "  Ep 01  train=x.xxx  val=x.xxx  (Xs)\n\n"
    "Meilleur val_acc : 0.xxxx\n"
    "Test accuracy    : 0.xxxx\n"
    "Modèle           : models/best_densenet121.pth\n"
    "MLflow run ID    : xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx\n\n"
    "MLflow Registry :\n"
    "  [OK] Modele blood-cell-densenet121 v1 enregistre\n"
    "  [OK] Alias @production -> v1  (premier modele)\n\n"
    "Résumé : val_acc=0.xxxx  test_acc=0.xxxx"
)
add_info_box(doc, "Avec 100 images et 1 epoch, les métriques seront faibles (~0.10-0.25). "
             "C'est attendu — ce test vérifie le pipeline, pas la qualité du modèle.")

doc.add_paragraph("Fichiers générés dans models/ :").runs[0].bold = True
add_code_block(doc,
    "models/\n"
    "  best_densenet121.pth         <- poids du meilleur modèle\n"
    "  confusion_matrix.png         <- matrice de confusion (8x8)\n"
    "  classification_report.txt    <- rapport précision/recall par classe\n"
    "  label_mapping.json           <- index -> nom de classe\n"
    "  metrics.json                 <- toutes les métriques du run\n"
    "  class_names.json             <- liste ordonnée des 8 classes"
)

# ── TEST 3 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 3 — Vérification MLflow UI", level=2)
add_warning_box(doc, "Nécessite que le Test 2 ait été lancé au moins une fois.")
add_info_box(doc, "Ce test vérifie manuellement que toutes les données sont bien loggées dans MLflow.")

p = doc.add_paragraph("Ouvrir : ")
r = p.add_run("http://localhost:5001")
r.font.name = "Courier New"
r.bold = True

doc.add_heading("3.1 — Expérience et run", level=3)
checks_exp = [
    ("Expérience visible", 'Expérience "bloodcells-densenet121" dans la liste de gauche'),
    ("Run listé", "Au moins 1 run avec statut FINISHED dans la liste"),
    ("Tags", 'Onglet Tags : git_commit=xxxxxxx, run_type=base'),
]
for check, detail in checks_exp:
    p = doc.add_paragraph()
    p.add_run(f"□  {check} : ").bold = True
    p.add_run(detail).font.size = Pt(10)

doc.add_heading("3.2 — Paramètres loggés", level=3)
add_code_block(doc,
    "batch_size      = 8\n"
    "lr_head         = 0.001\n"
    "lr_full         = 0.0001\n"
    "epochs_head     = 1\n"
    "epochs_full     = 1\n"
    "n_train         = 69\n"
    "n_val           = 16\n"
    "n_test          = 15\n"
    "optimizer       = AdamW\n"
    "part            = 0"
)

doc.add_heading("3.3 — Métriques loggées", level=3)
add_code_block(doc,
    "best_val_acc        = 0.xxxx\n"
    "test_acc            = 0.xxxx\n"
    "macro_f1            = 0.xxxx\n"
    "weighted_f1         = 0.xxxx\n"
    "precision_macro     = 0.xxxx\n"
    "recall_macro        = 0.xxxx\n"
    "recall_erythroblast = 0.xxxx   <- classe critique\n"
    "recall_ig           = 0.xxxx   <- classe critique\n"
    "train_time_s        = xxx.x\n"
    "n_params            = 7978568\n"
    "train_loss_ep1      = x.xxx    <- courbe par epoch\n"
    "val_loss_ep1        = x.xxx"
)

doc.add_heading("3.4 — Artifacts loggés", level=3)
checks_art = [
    "confusion_matrix.png — matrice 8×8 visualisable dans l'UI",
    "classification_report.txt — rapport par classe",
    "label_mapping.json — index → nom de classe",
    "model/ — modèle PyTorch enregistré (MLflow model format)",
]
for check in checks_art:
    doc.add_paragraph(f"□  {check}", style="List Bullet")

doc.add_heading("3.5 — Model Registry", level=3)
p = doc.add_paragraph("Aller dans l'onglet ")
r = p.add_run("Models")
r.bold = True
p.add_run(" dans la navigation MLflow.")

checks_reg = [
    ('Modèle "blood-cell-densenet121" présent dans la liste'),
    ("Version v1 (ou plus) avec statut None (MLflow 3.x, pas de stages)"),
    ("Alias @production assigné à la meilleure version"),
    ("Alias @challenger si un 2e run a été lancé et le modèle n'a pas dépassé la production"),
]
for check in checks_reg:
    doc.add_paragraph(f"□  {check}", style="List Bullet")

add_info_box(doc,
    "Règles de promotion : @production si macro_f1 >= prod ET recall_erythroblast >= prod - 0.02 "
    "ET recall_ig >= prod - 0.02. Sinon @challenger. Premier run → toujours @production.")

# ── TEST 4 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 4 — Prédiction via MLflow Registry + logging Supabase", level=2)
add_warning_box(doc, "Le serveur MLflow (Test 1) doit être démarré — predict_model.py "
                "charge le modèle @production depuis le Registry (plus de --model requis).")
add_info_box(doc, "Phase 2 : predict_model.py charge automatiquement le modèle @production "
             "depuis le MLflow Registry, logue inference_ms et enregistre dans Supabase.")

doc.add_heading("Terminal Mac", level=3)
add_code_block(doc,
    "# Mode Registry (défaut Phase 2 — charge @production automatiquement)\n"
    "python -m src.models.predict_model \\\n"
    "  --image data/Source_100/neutrophil/BNE_100878.jpg\n\n"
    "# Mode fallback .pth local (si MLflow indisponible)\n"
    "python -m src.models.predict_model \\\n"
    "  --image data/Source_100/neutrophil/BNE_100878.jpg \\\n"
    "  --model models/best_densenet121.pth"
)
doc.add_heading("Terminal Windows (PowerShell)", level=3)
add_code_block(doc,
    "# Mode Registry (défaut Phase 2 — charge @production automatiquement)\n"
    "python -m src.models.predict_model `\n"
    "  --image data/Source_100/neutrophil/BNE_100878.jpg\n\n"
    "# Mode fallback .pth local (si MLflow indisponible)\n"
    "python -m src.models.predict_model `\n"
    "  --image data/Source_100/neutrophil/BNE_100878.jpg `\n"
    "  --model models/best_densenet121.pth"
)

doc.add_paragraph("Résultat attendu (console) :").runs[0].bold = True
add_code_block(doc,
    "  [MLflow] Modele charge : models:/blood-cell-densenet121@production\n\n"
    "Prediction : NEUTROPHIL  (ou autre classe avec modèle test 1 epoch)\n"
    "Confiance  : xx.x%\n"
    "Inference  : xx.x ms    <- NOUVEAU en Phase 2\n\n"
    "Top 3 :\n"
    "  neutrophil      xx.x%\n"
    "  lymphocyte      xx.x%\n"
    "  monocyte        xx.x%"
)
add_info_box(doc, "Si une prédiction est erythroblast ou ig, la mention [CRITIQUE] apparaît. "
             "Si MLflow est indisponible, passer --model models/best_densenet121.pth en fallback.")

doc.add_paragraph("Vérification Supabase (DBeaver ou psql) :").runs[0].bold = True
add_code_block(doc,
    "SELECT image_name, predicted_class, confidence, created_at\n"
    "FROM predictions\n"
    "ORDER BY created_at DESC\n"
    "LIMIT 5;"
)
doc.add_paragraph("Résultat attendu :").runs[0].bold = True
add_code_block(doc,
    " image_name          | predicted_class | confidence | created_at\n"
    "---------------------+-----------------+------------+----------------------------\n"
    " BNE_100878.jpg      | neutrophil      |   0.xxxx   | 2026-06-17 xx:xx:xx"
)
add_info_box(doc, "Si Supabase est indisponible, le script continue sans erreur "
             "(mode silencieux) — seul un [warn] apparaît dans la console.")

# ── TEST 5 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 5 — API FastAPI avec MLflow intégré", level=2)
add_warning_box(doc, "Le serveur MLflow (Test 1) doit être démarré avant ce test.")

doc.add_heading("Étape 5.1 — Démarrer l'API", level=3)
doc.add_heading("Terminal Mac", level=3)
add_code_block(doc, "source .venv/bin/activate\nuvicorn src.serving.api:app --reload --port 8000")
doc.add_heading("Terminal Windows (PowerShell)", level=3)
add_code_block(doc, ".venv\\Scripts\\Activate.ps1\nuvicorn src.serving.api:app --reload --port 8000")

doc.add_heading("Étape 5.2 — Swagger UI (http://localhost:8000/docs)", level=3)
tests_swagger = [
    ("POST /training (groupe Entraînement)",
     'Try it out → body :\n{"data_dir":"data/Source_100","epochs_head":1,"epochs_full":1,"batch_size":8}\n→ Execute (~3-5 min)',
     '{\n  "status": "ok",\n  "val_acc": 0.xxxx,\n  "test_acc": 0.xxxx,\n'
     '  "mlflow_run_id": "xxxxxxxx",\n  "model_path": "models/best_densenet121.pth"\n}'),
    ("POST /predict (groupe Inférence)",
     "Try it out → Choose File → image dans data/Source_100/ → Execute",
     '{\n  "predicted_class": "neutrophil",\n  "confidence": 0.xxxx,\n'
     '  "is_critical": false,\n  "inference_ms": xx.x,\n  "top3": [...]\n}'),
]
for name, action, response in tests_swagger:
    doc.add_paragraph(name).runs[0].bold = True
    doc.add_paragraph(action)
    doc.add_paragraph("Réponse attendue :").runs[0].italic = True
    add_code_block(doc, response)

doc.add_heading("Étape 5.3 — Vérifier le run dans MLflow UI", level=3)
doc.add_paragraph(
    "Après le POST /training via Swagger, ouvrir http://localhost:5001 → "
    "un nouveau run doit apparaître dans l'expérience 'bloodcells-densenet121'."
)

# ── TEST 6 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 6 — Démonstration du garde-fou de promotion", level=2)
add_info_box(doc, "Ce test prouve qu'un modèle dégradé ne peut pas écraser @production, "
             "même s'il est plus récent. Le script simule un challenger avec recall_ig -30%.")

doc.add_heading("Terminal Mac & Windows (PowerShell)", level=3)
add_code_block(doc,
    "# Mac\n$env:PYTHONIOENCODING=utf-8  # Windows uniquement\n\n"
    "python scripts/demo_garde_fou.py"
)

doc.add_paragraph("Résultat attendu :").runs[0].bold = True
add_code_block(doc,
    "============================================================\n"
    "  Demo garde-fou de promotion MLflow\n"
    "============================================================\n\n"
    "[production] Version N\n"
    "  macro_f1             = 0.1894\n"
    "  recall_erythroblast  = 0.0000\n"
    "  recall_ig            = 0.5000\n\n"
    "[challenger simulé] Métriques dégradées artificiellement\n"
    "  macro_f1             = 0.1394  (prod - 0.05)\n"
    "  recall_ig            = 0.2000  (prod - 0.30)\n\n"
    "  macro_f1            : 0.1394 >= 0.1894  -> KO\n"
    "  recall_ig           : 0.2000 >= 0.4800  -> KO\n\n"
    "  [KO] GARDE-FOU ACTIVE — reste @challenger\n"
    "       Raison : macro_f1 0.1394 < 0.1894\n"
    "       Raison : recall_ig 0.2000 < 0.5000 - 0.02\n\n"
    "  @production reste sur Version N"
)

doc.add_paragraph("Vérification dans MLflow UI :").runs[0].bold = True
doc.add_paragraph(
    "Ouvrir http://localhost:5001 → Model registry → blood-cell-densenet121\n"
    "La version N conserve @production. La version N-1 affiche @challenger."
)
add_warning_box(doc, "Sur Windows : lancer avec $env:PYTHONIOENCODING='utf-8' "
                "pour éviter l'erreur d'encodage des emojis MLflow.")

doc.add_heading("Alternative — Déclencher @challenger via un vrai run", level=3)
doc.add_paragraph(
    "Il est aussi possible de faire apparaître @challenger en lançant un nouvel entraînement "
    "dont les métriques sont inférieures à la version @production. "
    "Le garde-fou bloque alors la promotion et le nouveau modèle reste @challenger."
)
add_warning_box(doc, "Avec 100 images et 1 epoch, le résultat est aléatoire : "
                "le run peut aussi passer les garde-fous et devenir @production (Version N+1 sans challenger). "
                "Utiliser --seed 999 augmente les chances d'obtenir de mauvaises métriques.")
add_code_block(doc,
    "# Windows (PowerShell)\n"
    "python -m src.train.training `\n"
    "  --data-dir data/Source_100 --output-dir models `\n"
    "  --epochs-head 1 --epochs-full 1 --batch-size 8 `\n"
    "  --run-type base --part 0 --seed 999\n\n"
    "# Mac\n"
    "python -m src.train.training \\\n"
    "  --data-dir data/Source_100 --output-dir models \\\n"
    "  --epochs-head 1 --epochs-full 1 --batch-size 8 \\\n"
    "  --run-type base --part 0 --seed 999"
)
add_info_box(doc, "Si le run échoue les garde-fous → Version N+1 reçoit @challenger dans le Registry. "
             "Si le run les passe → Version N+1 devient @production (démo moins parlante). "
             "Pour une démo garantie, préférer scripts/demo_garde_fou.py (résultat certain, ~5 sec).")

# ── TEST 7 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 7 — MLproject (reproductibilité, optionnel)", level=2)
add_info_box(doc, "MLproject permet de lancer le training avec mlflow run pour une reproductibilité totale. "
             "Utiliser --env-manager=local pour éviter la dépendance à conda.")

doc.add_heading("Terminal Mac & Windows (PowerShell)", level=3)
add_code_block(doc,
    "mlflow run . -e train --env-manager=local `\n"
    "  -P data_dir=data/Source_100 `\n"
    "  -P epochs_head=1 `\n"
    "  -P epochs_full=1 `\n"
    "  -P batch_size=8 `\n"
    "  -P part=0 `\n"
    "  -P run_type=base"
)
doc.add_paragraph("Résultat attendu :").runs[0].bold = True
add_code_block(doc,
    "2026/06/17 xx:xx:xx INFO mlflow.projects.utils: === Created directory ... ===\n"
    "2026/06/17 xx:xx:xx INFO mlflow.projects: === Running command 'python -m src.train.training ...' ===\n"
    "...\n"
    "2026/06/17 xx:xx:xx INFO mlflow.projects: === Run (ID 'xxxxxxxx') succeeded ==="
)
add_warning_box(doc, "MLFLOW_TRACKING_URI doit être exporté en variable d'environnement "
                "pour que mlflow run utilise le bon serveur :\n"
                "  export MLFLOW_TRACKING_URI=http://localhost:5001  (Mac)\n"
                "  $env:MLFLOW_TRACKING_URI='http://localhost:5001'  (Windows PowerShell)")

# ── TABLEAU RÉCAP ─────────────────────────────────────────────────────────────
doc.add_heading("Récapitulatif des tests Phase 2", level=2)
recap_rows = [
    ("MLflow serveur",      "docker-compose up -d --build mlflow",              "Container blood_cell_mlflow UP + UI http://localhost:5001"),
    ("Training MLflow",     "python -m src.train.training ... --run-type base", "run ID affiché + alias @production assigné"),
    ("MLflow UI — Params",  "http://localhost:5001 → run → Params",             "10 paramètres présents (batch_size, lr_head, n_train...)"),
    ("MLflow UI — Metrics", "http://localhost:5001 → run → Metrics",            "8 métriques dont recall_erythroblast et recall_ig"),
    ("MLflow UI — Artifacts","http://localhost:5001 → run → Artifacts",         "confusion_matrix.png + classification_report.txt visibles"),
    ("MLflow Registry",     "http://localhost:5001 → Model registry",           "blood-cell-densenet121 vN avec alias @production"),
    ("Predict Registry",    "python -m src.models.predict_model --image ...",   "[MLflow] Modele charge @production + inference_ms affiché"),
    ("Predict Supabase",    "Supabase → SELECT * FROM predictions",             "Nouvelle ligne avec image_name + confidence + created_at"),
    ("Garde-fou démo",      "python scripts/demo_garde_fou.py",                 "[KO] GARDE-FOU ACTIVE + @production inchangé dans Registry"),
    ("API /training",       "Swagger → POST /training",                         "mlflow_run_id présent dans la réponse JSON"),
    ("API /predict",        "Swagger → POST /predict",                          "inference_ms présent dans la réponse JSON"),
    ("MLproject (opt.)",    "mlflow run . -e train --env-manager=local",        "Run succeeded dans les logs mlflow"),
]
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdrs = ["Test", "Commande / Action", "Succès si..."]
for i, (cell, hdr) in enumerate(zip(table.rows[0].cells, hdrs)):
    set_cell_bg(cell, "1F497D")
    set_cell_border(cell)
    run = cell.paragraphs[0].add_run(hdr)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

for i, row_data in enumerate(recap_rows):
    row = table.add_row().cells
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    for j, (cell, content) in enumerate(zip(row, row_data)):
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        run = cell.paragraphs[0].add_run(content)
        run.font.size = Pt(9)
        if j == 2:
            run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)

doc.save(OUTPUT)
print(f"Guide généré : {OUTPUT}")
