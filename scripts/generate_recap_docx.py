"""Génère reports/recap_session.docx — récap CI fix + auth Streamlit + infra partagée."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parents[1]
OUT = ROOT / "reports" / "recap_session.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(18)
    return p


def main():
    doc = Document()

    title = doc.add_heading("Récap de session — CI, retraining & authentification Streamlit", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Projet : fev26_bmle_blood_cells — {date.today():%d/%m/%Y}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 1. Contexte
    add_heading(doc, "1. Contexte", level=1)
    doc.add_paragraph(
        "La CI GitHub Actions échouait sur le job \"Tests unitaires\" (flake8 puis pytest), "
        "bloquant le job \"Build Docker image\" qui en dépend. L'investigation a révélé deux "
        "problèmes distincts : des erreurs de lint accumulées dans le dépôt, et une "
        "incompatibilité de fond entre le modèle déployé et le dataset actuel."
    )

    # 2. CI / Lint
    add_heading(doc, "2. Correction de la CI (lint)", level=1)
    add_bullet(doc, "311 violations flake8 corrigées sur l'ensemble du dépôt (choix : tout corriger, "
               "pas de relâchement des règles CI).")
    add_bullet(doc, "Corrections automatiques via autopep8 (whitespace, lignes trop longues) "
               "puis corrections manuelles : imports inutilisés, noms de variables ambigus "
               "(self.l → self.labels), découpage de chaînes trop longues.")
    add_bullet(doc, "setup.cfg ajouté par Sara (en parallèle) pour ignorer certaines règles de "
               "style ML (E203, E221, E272) — fusionné sans conflit de fond.")
    add_bullet(doc, "flake8 src/ --max-line-length=120 (commande exacte de la CI) : 0 erreur.")

    # 3. Mismatch modèle / dataset
    add_heading(doc, "3. Incompatibilité modèle / dataset (9 vs 8 classes)", level=1)
    doc.add_paragraph(
        "Le test test_predict_returns_8_classes a révélé que best_DenseNet_121.pth (et "
        "best_ConvNeXt_Tiny.pth) ont une tête de sortie à 9 classes, alors que le dataset "
        "actuel n'en compte que 8. Cause racine : un script d'entraînement historique "
        "(Fred_DL_pipeline_full_v1.py) calculait dynamiquement la liste des classes à partir "
        "des sous-dossiers présents sur disque à l'époque, sans respecter son propre filtre "
        "EXPECTED_CLASSES. La 9e classe n'existe plus nulle part — non récupérable."
    )
    add_bullet(doc, "Un pull a d'abord été fait pour vérifier si une correction existait déjà "
               "en amont : Sara avait ajouté une détection dynamique du nombre de classes côté "
               "API (mitigation, évite le crash, mais ne corrige pas le modèle lui-même).")
    add_bullet(doc, "Décision : relancer l'entraînement via le nouveau pipeline MLflow "
               "training.py (python -m src.train.training) pour produire un checkpoint "
               "réellement compatible 8 classes. Entraînement en cours en arrière-plan "
               "au moment de la rédaction de ce récap (densenet121, device MPS).")

    # 4. Authentification Streamlit
    add_heading(doc, "4. Authentification Streamlit (nouvelle fonctionnalité)", level=1)
    doc.add_paragraph(
        "Mise en place d'une gestion de comptes pour restreindre l'accès à l'application "
        "Streamlit : table SQL côté Supabase, mots de passe hashés (jamais stockés en clair)."
    )

    add_heading(doc, "4.1 Choix techniques", level=2)
    add_bullet(doc, "Hashage bcrypt (bcrypt.hashpw / bcrypt.checkpw) plutôt qu'un chiffrement "
               "réversible : un hash ne peut pas être \"déchiffré\" pour retrouver le mot de "
               "passe, ce qui est la pratique standard pour des credentials.")
    add_bullet(doc, "Table users dans la base Supabase (PostgreSQL) déjà utilisée par le "
               "projet pour la table predictions — pas de nouvelle infra à provisionner.")

    add_heading(doc, "4.2 Fichiers créés", level=2)
    add_code(doc, "scripts/sql/001_create_users.sql   — schéma de la table users")
    add_code(doc, "src/auth/db.py                     — connexion Supabase (psycopg2)")
    add_code(doc, "src/auth/users.py                  — create_user / verify_user / "
                  "list_usernames / delete_user")
    add_code(doc, "src/auth/__init__.py                — exports du module")
    add_code(doc, "scripts/manage_users.py            — CLI admin (add / list / delete)")

    add_heading(doc, "4.3 Fichiers modifiés", level=2)
    add_bullet(doc, "requirements.txt : ajout de bcrypt>=4.0.")
    add_bullet(doc, "src/serving/app.py : ajout d'un écran de connexion (login_screen()) qui "
               "bloque l'accès à l'interface de classification tant que l'identifiant/mot de "
               "passe ne sont pas validés via verify_user(). Bouton de déconnexion ajouté "
               "dans la sidebar.")

    add_heading(doc, "4.4 Validation effectuée", level=2)
    add_bullet(doc, "Migration SQL exécutée avec succès sur la base Supabase réelle "
               "(table users créée).")
    add_bullet(doc, "Cycle complet testé en conditions réelles : création, vérification "
               "(mot de passe correct / incorrect / utilisateur inconnu), listing, "
               "suppression — via les fonctions Python et via le CLI manage_users.py.")
    add_bullet(doc, "flake8 sur tous les nouveaux fichiers : 0 erreur.")
    add_bullet(doc, "Suite pytest (hors tests dépendants du modèle en cours de "
               "réentraînement) : 8/8 tests passent.")
    add_bullet(doc, "Comptes de test supprimés après vérification — la table users est "
               "donc actuellement vide, ce qui est attendu : aucun compte réel n'a encore "
               "été créé.")

    # 5. Infrastructure partagée
    add_heading(doc, "5. Infrastructure partagée (Supabase + DagsHub)", level=1)
    doc.add_paragraph(
        "Mise en place et test des deux briques d'infrastructure partagées par l'équipe : "
        "Supabase (PostgreSQL, pour MLflow / résultats / comptes) et DagsHub (datalake DVC "
        "pour les images et les modèles)."
    )

    add_heading(doc, "5.1 Supabase — tables existantes", level=2)
    add_bullet(doc, "users — comptes Streamlit (id, username, password_hash bcrypt, created_at).")
    add_bullet(doc, "predictions — historique des prédictions (image_name, predicted_class, "
               "confidence, mlflow_run_id, triggered_by).")
    add_bullet(doc, "training_runs — logs des entraînements déclenchés via l'API "
               "(val_acc, test_acc, status...).")
    add_bullet(doc, "dataset_images — table créée, vide pour le moment.")
    add_bullet(doc, "TestCedric — table de test de connexion.")

    add_heading(doc, "5.2 DagsHub — fichiers versionnés (DVC)", level=2)
    add_bullet(doc, "models.dvc — best_DenseNet_121.pth (27 MB).")
    add_bullet(doc, "data/Source_100.dvc — 100 images de test, 8 classes (~12-13/classe).")
    add_bullet(doc, "data/Source_full.dvc — dataset complet data/raw, 17093 images.")
    add_bullet(doc, "data/runs/run1.dvc à run5.dvc — 5 runs d'acquisition simulés "
               "(800 images chacun, 100/classe x 8 classes), déjà sur DagsHub.")

    add_heading(doc, "5.3 Bascule vers le pooler Supabase (IPv4)", level=2)
    doc.add_paragraph(
        "L'hôte direct Supabase (db.<projet>.supabase.co) ne résout qu'en IPv6 sur le plan "
        "gratuit. Docker Desktop (Mac) ne route pas l'IPv6, et certains réseaux/Windows non "
        "plus selon la configuration. Pour fiabiliser la connexion pour toute l'équipe, "
        "SUPABASE_HOST a été basculé vers le pooler Supavisor en mode session "
        "(aws-0-eu-west-1.pooler.supabase.com, IPv4), avec SUPABASE_USER au format "
        "postgres.<id_projet>. Fonctionne désormais depuis Docker, Mac et Windows, "
        "indépendamment du support IPv6 du réseau."
    )

    add_heading(doc, "5.4 Scripts de test créés", level=2)
    add_code(doc, "scripts/test_connections.py   — teste Supabase (insert/read/delete) et "
                  "DagsHub (manifests DVC)")
    add_code(doc, "scripts/print_test_cedric.py  — affiche la table TestCedric + compte les "
                  "images du datalake")
    add_code(doc, "Dockerfile.test_db            — image Docker légère pour faire tourner ce "
                  "test sans installer Python localement")

    # 6. Nettoyage des chemins personnels
    add_heading(doc, "6. Suppression des chemins personnels codés en dur", level=1)
    doc.add_paragraph(
        "Plusieurs scripts contenaient des chemins absolus propres à la machine de leur "
        "auteur (Fred, Romane, ou Sara sous Windows), ce qui provoquait des allers-retours "
        "à chaque commit (chacun remettait son propre chemin). Ces chemins ont été déplacés "
        "vers des variables d'environnement, lues depuis .env (fichier jamais commité)."
    )
    add_bullet(doc, "scripts/find_best_tiffs.py, validate_on_cancer_archive.py, "
               "compare_1fold_vs_5fold.py, src/train/train_simple.py — chemins de caches "
               "personnels (CancerImagingArchive, OneDrive, Mendeley brut) déplacés vers "
               "CROSSVAL_CACHE_DIR / CANCER_ARCHIVE_DIR / ONEDRIVE_CACHE_DIR / MENDELEY_RAW_DIR.")
    add_bullet(doc, "src/evaluation/eval_experiments.py, eval_best_models.py — chemin Windows "
               "de Sara (dataset Acevedo) déplacé vers ACEVEDO_DATA_DIR ; le dossier de sortie "
               "est maintenant relatif au dépôt (reports/Sara_DL_convnext_densenet_hyperparam).")
    add_bullet(doc, "scripts/generate_report.py — chemin de sortie de Romane rendu relatif "
               "au dépôt.")
    add_bullet(doc, "src/serving/batch_inference.py — dossier de démo par défaut surchargeable "
               "via DEMO_IMAGES_DIR, avec repli automatique sur data/Source_100 (partagé).")
    add_bullet(doc, "bcrypt ajouté à requirements/base.txt et requirements/streamlit.txt "
               "(manquant, nécessaire pour l'authentification).")
    doc.add_paragraph(
        "Chaque script lève désormais une erreur claire si la variable d'environnement "
        "attendue n'est pas définie, au lieu d'échouer silencieusement ou de pointer vers "
        "le mauvais dossier."
    )

    # 7. Configuration VSCode / .env pour Sara et Romane
    add_heading(doc, "7. À faire par Sara et Romane : configurer son .env et VSCode", level=1)
    doc.add_paragraph(
        "Pour que les scripts fonctionnent sans toucher au code, chacune doit créer son "
        "propre fichier .env local (jamais commité — il est dans .gitignore)."
    )

    add_heading(doc, "Étape 1 — Récupérer les derniers changements", level=2)
    add_code(doc, "git pull origin main")

    add_heading(doc, "Étape 2 — Créer son .env personnel", level=2)
    add_code(doc, "cp .env.example .env          # Mac/Linux")
    add_code(doc, "copy .env.example .env         # Windows (cmd)")
    doc.add_paragraph(
        "Remplir dans .env : DAGSHUB_TOKEN (token personnel, dagshub.com > Settings > "
        "Tokens) et SUPABASE_PASSWORD (à demander à l'équipe). Les autres valeurs "
        "(SUPABASE_HOST, DAGSHUB_REPO...) sont déjà correctes dans .env.example."
    )
    doc.add_paragraph(
        "Les variables de chemins personnels (CROSSVAL_CACHE_DIR, CANCER_ARCHIVE_DIR, "
        "ACEVEDO_DATA_DIR...) sont commentées dans .env.example : à décommenter et remplir "
        "UNIQUEMENT si vous lancez les scripts de recherche qui en ont besoin "
        "(eval_experiments.py / eval_best_models.py pour Sara, par exemple). Inutile pour "
        "Streamlit, l'API ou l'authentification."
    )

    add_heading(doc, "Étape 3 — Créer/activer l'environnement virtuel", level=2)
    add_code(doc, "python -m venv .venv")
    add_code(doc, "source .venv/bin/activate      # Mac/Linux")
    add_code(doc, ".venv\\Scripts\\activate          # Windows")
    add_code(doc, "pip install -r requirements.txt")

    add_heading(doc, "Étape 4 — Sélectionner l'interpréteur dans VSCode", level=2)
    add_bullet(doc, "Ouvrir la palette de commandes : Cmd+Shift+P (Mac) ou Ctrl+Shift+P "
               "(Windows).")
    add_bullet(doc, "Taper \"Python: Select Interpreter\".")
    add_bullet(doc, "Choisir celui du dossier .venv du projet (et non un Python global).")

    add_heading(doc, "Étape 5 — Vérifier que tout fonctionne", level=2)
    add_code(doc, "python scripts/test_connections.py")
    doc.add_paragraph(
        "Doit afficher [OK] pour Supabase et DagsHub. Puis pour tester l'authentification : "
    )
    add_code(doc, "streamlit run src/serving/app.py")
    doc.add_paragraph(
        "Se connecter avec son identifiant (fred / sara / romane) — demander le mot de "
        "passe à la personne concernée si besoin, ou utiliser "
        "python scripts/manage_users.py add <identifiant> pour créer son propre compte."
    )

    # 8. État actuel / suite
    add_heading(doc, "8. État actuel et prochaines étapes", level=1)
    add_bullet(doc, "La table users contient désormais 3 comptes (fred, sara, romane) — "
               "l'authentification Streamlit est opérationnelle.")
    add_bullet(doc, "Entraînement en arrière-plan à surveiller : une fois terminé, vérifier "
               "que models/best_densenet121.pth est bien compatible 8 classes et que "
               "test_predict_returns_8_classes passe.")
    add_bullet(doc, "Sara et Romane doivent configurer leur .env local (voir section 7) "
               "avant de pull/relancer le projet.")
    add_bullet(doc, "Commit/push des changements (lint CI, auth Streamlit, infra, chemins) "
               "déjà effectué sur main.")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
