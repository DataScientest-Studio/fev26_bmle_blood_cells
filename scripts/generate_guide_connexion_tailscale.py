"""Génère reports/Fred/guide_connexion_tailscale.docx — guide pratique pour
Romane et Sara : comment se connecter au Mac de Fred (Airflow + MLflow) et
lancer un nouvel entraînement.

Contient des IP Tailscale réelles, lues depuis .env. Document PRIVÉ, à
transmettre directement à Romane et Sara (email, Slack...) — jamais sur
GitHub (repo public).
"""

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

ROOT = Path(__file__).parents[1]
load_dotenv(ROOT / ".env")
OUT = ROOT / "reports" / "Fred" / "guide_connexion_tailscale.docx"

if not os.getenv("MAC_TAILSCALE_IP"):
    raise EnvironmentError("MAC_TAILSCALE_IP doit être défini dans ton .env local.")
MAC_IP = os.environ["MAC_TAILSCALE_IP"]


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(18)
    return p


def main():
    doc = Document()

    title = doc.add_heading("Guide de connexion — Airflow & MLflow (Mac de Fred)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Pour Romane et Sara — fev26_bmle_blood_cells — {date.today():%d/%m/%Y}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Document privé — ne pas publier (contient des adresses réseau personnelles).")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0x33, 0x33)

    doc.add_paragraph(
        "L'entraînement du modèle tourne maintenant sur le PC Windows de Fred (qui a une "
        "carte GPU adaptée), mais tout est piloté depuis Airflow, installé sur son Mac. Pour "
        "suivre les entraînements ou en lancer un nouveau, il suffit d'accéder aux interfaces "
        "web Airflow et MLflow sur son Mac — pas besoin d'installer quoi que ce soit d'autre "
        "ni d'avoir un accès terminal."
    )

    # Étape 1
    add_heading(doc, "Étape 1 — Installer Tailscale et recevoir l'accès", level=1)
    doc.add_paragraph(
        "Fred va te partager uniquement son Mac via la fonction \"Share device\" de Tailscale "
        "(pas une invitation à son réseau complet) : tu auras accès à cet appareil précis, et "
        "rien d'autre sur son réseau."
    )
    add_bullet(doc, "Installer l'application Tailscale : tailscale.com/download (Mac, Windows, "
               "ou mobile selon ce que tu utilises).")
    add_bullet(doc, "Créer un compte Tailscale si tu n'en as pas déjà un (connexion possible "
               "avec Google/Microsoft/GitHub).")
    add_bullet(doc, "Accepter l'invitation de partage envoyée par Fred (lien reçu par email).")
    add_bullet(doc, "Une fois connectée, le Mac de Fred doit apparaître dans ta liste "
               "d'appareils Tailscale.")

    # Étape 2
    add_heading(doc, "Étape 2 — Ouvrir Airflow (suivre / déclencher les entraînements)", level=1)
    add_code(doc, f"http://{MAC_IP}:8080")
    doc.add_paragraph("Identifiants : à demander à Fred (pas admin/admin par défaut).")
    add_bullet(doc, "Dans la liste des DAGs, ouvrir blood_cell_training_pipeline.")
    add_bullet(doc, "Vue \"Grid\" : historique des exécutions, statut de chaque tâche (verte = "
               "réussie, rouge = échouée, jaune/en cours = en cours).")
    add_bullet(doc, "Cliquer sur une tâche puis \"Logs\" pour voir le détail — la tâche "
               "train_model affiche la progression de l'entraînement en direct (epochs, "
               "accuracy...).")

    # Étape 3
    add_heading(doc, "Étape 3 — Lancer un nouvel entraînement", level=1)
    add_bullet(doc, "Dans Airflow, ouvrir le DAG blood_cell_training_pipeline.")
    add_bullet(doc, "Cliquer sur le bouton ▶ (\"Trigger DAG\") en haut à droite de l'écran.")
    add_bullet(doc, "L'entraînement se lance automatiquement sur le PC Windows de Fred — il "
               "faut que ce PC soit allumé et connecté à Internet (Tailscale) pour que ça "
               "fonctionne, sinon la tâche train_model échoue.")
    add_bullet(doc, "Suivre la progression via les logs de train_model. Le résultat final "
               "(modèle promu en production, ou pas) apparaît ensuite dans la tâche "
               "check_promotion.")
    doc.add_paragraph(
        "À savoir : un entraînement complet prend environ 1h30 à 2h. Le PC Windows étant "
        "aussi utilisé pour jouer, éviter de lancer un entraînement en même temps qu'une "
        "session de jeu (les deux se ralentissent mutuellement, le GPU ne peut pas vraiment "
        "faire les deux en parallèle)."
    )

    # Étape 4
    add_heading(doc, "Étape 4 — Ouvrir MLflow (voir les résultats des modèles)", level=1)
    add_code(doc, f"http://{MAC_IP}:5001")
    add_bullet(doc, "Onglet \"Models\" → blood-cell-densenet121 : liste des versions, avec un "
               "tag generation (v0 = ancienne référence, v1/v2... = nouveaux cycles) et un "
               "alias @production sur la version actuellement utilisée.")
    add_bullet(doc, "Onglet \"Experiments\" → blood_cell_crossval_ameliorees : détail des "
               "courbes d'entraînement (loss/accuracy) par fold.")

    # Aide
    add_heading(doc, "En cas de problème", level=1)
    add_bullet(doc, "Page blanche / inaccessible : vérifier que Tailscale est bien connecté "
               "(icône dans la barre de menu/système) et que le Mac de Fred est allumé.")
    add_bullet(doc, "La tâche train_model échoue immédiatement : le PC Windows de Fred est "
               "probablement éteint ou déconnecté de Tailscale — lui demander de vérifier.")
    add_bullet(doc, "Pour toute autre question, contacter Fred directement.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
