"""Génère src/evidently/rapport_alertes_email_drift.docx —
rapport détaillé sur le système d'alertes email pour le monitoring drift
(IVDR 2017/746), incluant le guide de mise en place.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parents[1]
OUT = ROOT / "src" / "evidently" / "rapport_alertes_email_drift.docx"

BLUE       = "1F497D"
ORANGE     = "C0504D"
LIGHT_BLUE = "EBF3FB"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows, header_color=BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, hdr in zip(table.rows[0].cells, headers):
        set_cell_bg(cell, header_color)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
    for i, values in enumerate(rows):
        row = table.add_row().cells
        bg = LIGHT_BLUE if i % 2 == 0 else "FFFFFF"
        for cell, content in zip(row, values):
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            r = cell.paragraphs[0].add_run(content)
            r.font.size = Pt(8.5)
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p.paragraph_format.left_indent = Pt(20)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def main():
    doc = Document()

    # ── Titre ─────────────────────────────────────────────────────────────────
    title = doc.add_heading(
        "Système d'alertes email — Monitoring drift Evidently", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Projet fev26_bmle_blood_cells — {date.today():%d/%m/%Y}"
        "  ·  Auteur : Romane Beaurepere"
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(10)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub2.add_run("Contexte réglementaire : IVDR 2017/746 — Art. 10 & MDCG 2020-1")
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    run2.font.size = Pt(9)

    doc.add_paragraph()

    # ── 1. Contexte ───────────────────────────────────────────────────────────
    add_heading(doc, "1. Contexte et objectif", level=1)
    doc.add_paragraph(
        "Le modèle Blood Cell Classifier est soumis aux obligations de surveillance "
        "post-marché de l'IVDR 2017/746. Cela implique de détecter rapidement toute "
        "dégradation de la qualité des données en entrée ou des performances du modèle "
        "en production."
    )
    doc.add_paragraph(
        "Le module Evidently (src/evidently/drift_report.py) calculait déjà les "
        "indicateurs de drift, mais les rapports n'étaient générés que manuellement "
        "via un bouton dans l'interface Streamlit. Ce système ne permettait pas de "
        "détecter un problème sans intervention humaine."
    )
    p = doc.add_paragraph()
    run = p.add_run("Objectif : ")
    run.bold = True
    p.add_run(
        "automatiser la génération du rapport de drift chaque nuit et envoyer une "
        "notification par email à l'équipe dès qu'un seuil d'alerte IVDR est franchi."
    )

    # ── 2. Architecture ───────────────────────────────────────────────────────
    add_heading(doc, "2. Architecture de la solution", level=1)
    doc.add_paragraph(
        "La solution repose sur deux nouveaux fichiers et une mise à jour de la "
        "configuration :"
    )
    add_table(
        doc,
        ["Fichier", "Rôle"],
        [
            ["src/monitoring/email_alert.py",
             "Module d'envoi email via SMTP (stdlib Python, sans dépendance externe). "
             "Construit l'email HTML et l'envoie si le niveau de drift dépasse le seuil."],
            ["airflow/dags/blood_cell_drift_monitoring.py",
             "DAG Airflow planifié chaque nuit à 7h. Orchestre les 3 tâches : "
             "génération du rapport, vérification des performances, envoi de l'alerte."],
            [".env.example",
             "5 nouvelles variables SMTP ajoutées avec les 3 adresses de destinataires "
             "pré-remplies."],
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "2.1 Flux d'exécution du DAG", level=2)
    doc.add_paragraph(
        "Les tâches 1 et 2 s'exécutent en parallèle, puis la tâche 3 attend "
        "leurs résultats :"
    )
    add_table(
        doc,
        ["Tâche", "Action", "Si échec"],
        [
            ["1 — generate_drift_report",
             "Appelle generate_report() : calcule data drift, prediction drift et "
             "model drift. Sauvegarde le rapport dans Supabase (table drift_reports).",
             "Le DAG échoue et notifie Airflow."],
            ["2 — check_performance",
             "Appelle load_performance_metrics() : compare macro_f1 et F1 des classes "
             "critiques entre générations. Détecte les baisses > 5% (IVDR).",
             "Le DAG échoue et notifie Airflow."],
            ["3 — send_alert_if_needed",
             "Envoie l'email d'alerte HTML si le niveau global dépasse le seuil "
             "configuré. Sinon, log \"aucune alerte\" et termine silencieusement.",
             "Le DAG échoue et notifie Airflow."],
        ],
    )

    # ── 3. Seuils ─────────────────────────────────────────────────────────────
    add_heading(doc, "3. Seuils d'alerte (IVDR / ISO 14971)", level=1)
    add_table(
        doc,
        ["Indicateur", "Warning", "Alerte", "Critique"],
        [
            ["Data drift score (share of drifted columns)", "> 0.10", "> 0.20", "> 0.30"],
            ["Prediction drift score", "—", "> 0.20", "—"],
            ["Désaccord médecin (feedback)", "> 10 %", "> 15 %", "—"],
            ["Baisse macro_f1 vs baseline", "—", "> 5 %", "—"],
            ["Baisse F1 Erythroblast ou IG", "—", "—", "> 5 %"],
        ],
        header_color=ORANGE,
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Le niveau global retenu pour l'email est le plus haut niveau parmi tous les "
        "indicateurs. Par défaut, un email est envoyé dès le niveau warning. "
        "Ce seuil minimum est configurable via le paramètre min_level de "
        "send_drift_alert() dans le DAG."
    )

    # ── 4. Email ──────────────────────────────────────────────────────────────
    add_heading(doc, "4. Contenu de l'email d'alerte", level=1)
    add_bullet(doc,
        "[ALERTE] Drift détecté — Blood Cell Classifier (2026-06-30 07:00 UTC)",
        bold_prefix="Objet : ")
    add_bullet(doc,
        "Tableau HTML des métriques avec badges colorés par niveau "
        "(NORMAL vert, WARNING jaune, ALERTE orange, CRITIQUE rouge)",
        bold_prefix="Corps : ")
    add_bullet(doc,
        "delabot.frederic@gmail.com, diomand.sara@gmail.com, beaurepere.romane@gmail.com",
        bold_prefix="Destinataires : ")

    doc.add_paragraph()
    doc.add_paragraph("Le tableau de l'email contient les informations suivantes :")
    add_table(
        doc,
        ["Ligne", "Valeur exemple"],
        [
            ["Data drift score", "0.213  →  ALERTE"],
            ["Features driftées", "3"],
            ["Prediction drift score", "0.087  →  NORMAL"],
            ["Model drift (désaccord médecin)", "0.120"],
            ["Images référence / courantes", "500 / 142"],
        ],
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Si des alertes de performance IVDR sont détectées (baisse macro_f1 ou F1 "
        "de classes critiques), elles apparaissent en rouge sous le tableau. "
        "Les nuits sans anomalie (score < 0.10), aucun email n'est envoyé — "
        "le rapport est quand même sauvegardé silencieusement en base."
    )

    # ── 5. Mise en place ──────────────────────────────────────────────────────
    add_heading(doc, "5. Guide de mise en place", level=1)

    add_heading(doc, "5.1 Qui doit configurer quoi ?", level=2)
    doc.add_paragraph(
        "Le système d'envoi email est centralisé sur la machine qui héberge Airflow "
        "(le Mac de Romane). Sara et Fred n'ont rien à configurer — ils recevront "
        "les alertes automatiquement sur leur boite Gmail."
    )
    add_table(
        doc,
        ["Membre", "Action requise"],
        [
            ["Romane (Mac — héberge Airflow)",
             "Configurer les variables SMTP dans .env + redémarrer Airflow. "
             "C'est la seule machine qui envoie les emails. ✅ Déjà fait."],
            ["Fred (PC Windows — GPU fine-tuning)",
             "Aucune configuration SMTP nécessaire. "
             "Recevoir les emails sur delabot.frederic@gmail.com."],
            ["Sara",
             "Aucune configuration SMTP nécessaire. "
             "Recevoir les emails sur diomand.sara@gmail.com."],
        ],
    )

    add_heading(doc, "5.2 Configuration sur la machine Airflow (Romane — déjà fait)", level=2)
    doc.add_paragraph(
        "Ces étapes ont été réalisées le 29 juin 2026 sur le Mac de Romane. "
        "Elles sont documentées ici pour référence ou en cas de réinstallation."
    )

    add_heading(doc, "5.2.1 Créer un App Password Gmail", level=2)
    doc.add_paragraph(
        "Google bloque les connexions SMTP directes depuis 2022. Il faut utiliser "
        "un mot de passe d'application (App Password) et non le mot de passe du compte. "
        "La validation en deux étapes doit être activée au préalable."
    )
    add_numbered(doc, "Connectez-vous sur myaccount.google.com → Sécurité")
    add_numbered(doc, "Activez la validation en deux étapes si ce n'est pas déjà fait")
    add_numbered(doc, "Dans le champ de recherche, tapez \"Mots de passe des applications\"")
    add_numbered(doc, "Créez un nouveau mot de passe, nommez-le BloodCell Airflow")
    add_numbered(doc,
        "Copiez les 16 caractères générés (4 blocs de 4 lettres) — "
        "ils ne s'affichent qu'une seule fois")
    p = doc.add_paragraph()
    r = p.add_run(
        "⚠  Point d'attention : Google peut demander une confirmation de sécurité "
        "avant d'activer le mot de passe. Valider cette confirmation dans Gmail "
        "avant de lancer le test SMTP."
    )
    r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x50, 0x4D)

    add_heading(doc, "5.2.2 Configurer les variables dans .env", level=2)
    doc.add_paragraph(
        "Ajoutez les lignes suivantes dans le fichier .env à la racine du projet "
        "(sur la machine hébergeant Airflow) :"
    )
    add_code_block(doc,
        "SMTP_HOST=smtp.gmail.com\n"
        "SMTP_PORT=587\n"
        "SMTP_USER=beaurepere.romane@gmail.com\n"
        "SMTP_PASSWORD=xxxx xxxx xxxx xxxx   # App Password Gmail (16 caractères)\n"
        "ALERT_EMAILS=delabot.frederic@gmail.com,diomand.sara@gmail.com,"
        "beaurepere.romane@gmail.com"
    )
    p = doc.add_paragraph()
    run = p.add_run("⚠  Ne committez jamais ces valeurs dans Git. ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x4D)
    p.add_run("Le fichier .env est dans .gitignore.")

    add_heading(doc, "5.2.3 Transmettre les variables à Airflow (Docker Compose)", level=2)
    doc.add_paragraph(
        "Dans airflow/docker-compose-airflow.yml, ajoutez les variables sous la clé "
        "environment du service airflow-scheduler (et airflow-worker si présent) :"
    )
    add_code_block(doc,
        "environment:\n"
        "  SMTP_HOST: ${SMTP_HOST}\n"
        "  SMTP_PORT: ${SMTP_PORT}\n"
        "  SMTP_USER: ${SMTP_USER}\n"
        "  SMTP_PASSWORD: ${SMTP_PASSWORD}\n"
        "  ALERT_EMAILS: ${ALERT_EMAILS}"
    )

    add_heading(doc, "5.2.4 Redémarrer Airflow", level=2)
    add_code_block(doc,
        "docker compose -f airflow/docker-compose-airflow.yml down\n"
        "docker compose -f airflow/docker-compose-airflow.yml up -d"
    )
    doc.add_paragraph(
        "Le DAG blood_cell_drift_monitoring apparaîtra automatiquement dans l'UI Airflow."
    )

    # ── 6. Test réalisé ───────────────────────────────────────────────────────
    add_heading(doc, "6. Test réalisé le 29 juin 2026", level=1)
    doc.add_paragraph(
        "Un test manuel a été exécuté directement depuis un terminal sur le Mac de Romane "
        "pour valider l'envoi SMTP avant la mise en production dans Airflow. "
        "Aucun fichier de test dédié n'a été créé — le test a utilisé un rapport de "
        "drift simulé (données fictives) injecté directement dans send_drift_alert()."
    )

    add_heading(doc, "6.1 Commande de test exécutée", level=2)
    add_code_block(doc,
        "# Depuis la racine du projet, avec le .venv activé\n"
        ".venv/bin/python -c \"\n"
        "from dotenv import load_dotenv; load_dotenv('.env')\n"
        "from src.monitoring.email_alert import send_drift_alert\n\n"
        "# Rapport simulé avec data_drift_score=0.15 (niveau warning)\n"
        "fake_result = {\n"
        "    'data_drift_score': 0.15, 'data_drift_level': 'warning',\n"
        "    'pred_drift_score': 0.05, 'pred_drift_level': 'normal',\n"
        "    'model_drift_score': 0.08, 'n_drifted_features': 2,\n"
        "    'n_reference': 500, 'n_current': 142, 'model_version': 'v100',\n"
        "    'metrics': {'prediction_drift': {'confidence_drift': 0.04},\n"
        "                'data_drift': {'per_feature': {}},\n"
        "                'model_drift': {'n_feedback': 12, 'accuracy': 0.92}}\n"
        "}\n"
        "send_drift_alert(fake_result, min_level='warning')\n"
        "\""
    )

    add_heading(doc, "6.2 Résultats", level=2)
    add_table(
        doc,
        ["Étape", "Résultat"],
        [
            ["Connexion SMTP (smtp.gmail.com:587)", "✅ OK"],
            ["Authentification App Password Gmail", "✅ OK (après confirmation sécurité Google)"],
            ["Construction email HTML", "✅ OK"],
            ["Envoi à beaurepere.romane@gmail.com", "✅ Email reçu en boite de réception"],
            ["Log terminal", "Email d'alerte [WARNING] envoyé à beaurepere.romane@gmail.com"],
        ],
        header_color="2E7D32",
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Point d'attention rencontré lors du test : le premier App Password saisi "
        "était incomplet (15 caractères au lieu de 16 — un caractère manquant à la copie). "
        "L'authentification a échoué jusqu'à saisie du mot de passe correct. "
        "Toujours vérifier que le mot de passe fait bien 4 blocs de 4 lettres (16 caractères)."
    )

    add_heading(doc, "6.3 Test via l'UI Airflow (à faire après redémarrage)", level=2)
    doc.add_paragraph(
        "Pour valider le DAG complet (avec vrai rapport Evidently depuis Supabase) :"
    )
    add_numbered(doc,
        "Dans l'UI Airflow, ouvrir le DAG blood_cell_drift_monitoring")
    add_numbered(doc,
        "Cliquer Trigger DAG (▶) pour déclencher manuellement")
    add_numbered(doc,
        "Vérifier les logs de la tâche send_alert_if_needed — succès attendu :")
    add_code_block(doc,
        "Email d'alerte [WARNING] envoyé à delabot.frederic@gmail.com, "
        "diomand.sara@gmail.com, beaurepere.romane@gmail.com"
    )
    add_numbered(doc,
        "Vérifier la réception sur les trois boites mail")

    # ── 7. Planification ──────────────────────────────────────────────────────
    add_heading(doc, "7. Planification des DAGs", level=1)
    add_table(
        doc,
        ["DAG", "Planification", "Heure"],
        [
            ["blood_cell_pipeline", "Chaque dimanche", "2h00 — entraînement complet"],
            ["blood_cell_incremental_finetune_pipeline",
             "Chaque dimanche", "4h00 — fine-tuning incrémental"],
            ["blood_cell_drift_monitoring",
             "Chaque jour", "7h00 — monitoring drift + alerte email"],
        ],
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Le DAG de monitoring est entièrement indépendant des deux DAGs d'entraînement. "
        "Il peut être désactivé ou modifié sans aucun impact sur la production du modèle."
    )

    # ── 8. Évolutions possibles ───────────────────────────────────────────────
    add_heading(doc, "8. Évolutions possibles", level=1)
    add_bullet(doc,
        " : passer min_level=\"warning\" à min_level=\"alerte\" dans le DAG "
        "pour réduire le nombre d'emails.",
        bold_prefix="Modifier le seuil d'envoi")
    add_bullet(doc,
        " : modifier ALERT_EMAILS dans .env (liste séparée par virgules).",
        bold_prefix="Ajouter des destinataires")
    add_bullet(doc,
        " : modifier schedule_interval dans le DAG "
        "(ex: \"0 7 * * 1\" pour hebdomadaire le lundi).",
        bold_prefix="Changer la fréquence")
    add_bullet(doc,
        " : changer SMTP_HOST et SMTP_PORT "
        "(ex: Outlook → smtp.office365.com:587).",
        bold_prefix="Utiliser un autre fournisseur SMTP")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
