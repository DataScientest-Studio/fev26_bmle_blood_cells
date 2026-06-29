"""Génère reports/Fred/recap_weekend_26_27juin2026.docx — récap des
modifications du vendredi 26 / samedi 27 juin, et rapport du test de bout
en bout du DAG d'incremental fine-tuning (29 juin, PC Windows rallume pour
l'occasion).

Document de suivi de projet, pas committé sur GitHub (cf. convention des
autres récaps dans reports/) — à partager avec l'équipe par un canal interne
si besoin.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parents[1]
OUT = ROOT / "reports" / "Fred" / "recap_weekend_26_27juin2026.docx"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows, header_color="1F497D"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, hdr in zip(table.rows[0].cells, headers):
        set_cell_bg(cell, header_color)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
    for i, values in enumerate(rows):
        row = table.add_row().cells
        bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        for cell, content in zip(row, values):
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            r = cell.paragraphs[0].add_run(content)
            r.font.size = Pt(8.5)
    return table


def main():
    doc = Document()

    title = doc.add_heading("Récap weekend 26-27 juin + test bout en bout 29 juin", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Projet fev26_bmle_blood_cells — {date.today():%d/%m/%Y}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Modifications — vendredi 26 juin", level=1)

    add_heading(doc, "1.1 Déjà poussées sur main", level=2)
    add_table(
        doc,
        ["Commit", "Auteur", "Contenu"],
        [
            ["abe17b9", "Sara", "Refonte Streamlit (sidebar, grille GradCAM++ batch) + endpoint /gradcam"],
            ["108e055", "Fred/Claude", "Indice de confiance par classe + flux de feedback médecin"],
            ["93327fc", "Romane", "Fix : MLflow @production en priorité sur le .pth local (load_model())"],
            ["2a525d9", "Fred/Claude", "Fix flake8 (E231, E127) suite au merge"],
            ["31b3715", "Fred/Claude", "Merge de main (redesign Sara + Evidently Romane) dans master"],
            ["0c41007", "Fred/Claude", "Recap des corrections du pptx de soutenance"],
        ],
    )

    add_heading(doc, "1.2 Hors-code", level=2)
    add_bullet(
        doc,
        " : 17 093→17 092 fichiers, libellé erroné run1-5.dvc \"lots TIFF\" "
        "(en réalité JPEG) corrigé, KPI vide \"~ ms\" → \"~42 ms\", suppression d'une diapo "
        "doublon (coquille \"Neutrophil\").",
        bold_prefix="Corrections pptx soutenance",
    )
    add_bullet(
        doc,
        " : mot de passe régénéré, .env mis à jour, port pooler "
        "corrigé 5432 → 6543, conteneurs api/streamlit/mlflow + Airflow recréés et "
        "connexion re-testée.",
        bold_prefix="Mot de passe Supabase",
    )
    add_bullet(
        doc,
        " : root cause identifiée (flake8, pas les tests unitaires) et corrigée — "
        "build/tests/lint repassent au vert.",
        bold_prefix="CI cassée",
    )

    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Modifications — samedi 27 juin (en attente de commit)", level=1)
    doc.add_paragraph(
        "Ces changements sont dans l'arbre de travail local (pas encore committés ni "
        "poussés) — à valider avant un prochain push."
    )

    add_table(
        doc,
        ["Fichier", "Changement"],
        [
            ["docker/api/Dockerfile", "Ajout de libgl1 (système) — /gradcam plantait "
             "systématiquement (cv2 → libGL.so.1 manquant), endpoint jamais "
             "fonctionnel jusqu'ici"],
            ["requirements/streamlit.txt", "Ajout d'evidently==0.7.21, absent depuis son "
             "intégration — l'onglet Monitoring n'avait jamais vraiment tourné en "
             "conteneur"],
            ["scripts/init_db.py", "Nouvelle colonne predictions.patient_id (INTEGER)"],
            ["src/evidently/drift_report.py", "list_confusion_generations() + "
             "load_confusion_matrix() — matrices de confusion par génération"],
            ["src/serving/api.py", "patient_id (form-data) accepté sur /predict et "
             "/gradcam, propagé jusqu'à Supabase"],
            ["src/serving/app.py", "Onglet Monitoring complété (matrice de confusion, "
             "détail du drift par feature, drift de confidence) + patient simulé "
             "(1 lot analysé = 1 patient) + nouvel onglet Recherche (par nom d'image "
             "et par patient)"],
            ["airflow/dags/_common.py", "next_generation() (calcul auto depuis le "
             "Registry MLflow) + next_tiff_batch() (suivi du prochain lot via "
             "Airflow Variable)"],
            ["airflow/dags/blood_cell_incremental_finetune_pipeline.py",
             "Planification ajoutée (dimanche 4h) + sélection automatique du lot/"
             "génération — déclenchement manuel toujours possible"],
        ],
    )

    add_heading(doc, "2.1 Données de test générées", level=2)
    add_bullet(
        doc,
        " (via /predict, mix réaliste des 8 classes par lot) pour "
        "valider le nouvel onglet Recherche.",
        bold_prefix="500 prédictions de test — 10 patients simulés x 50 images",
    )

    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Rapport de test — DAG incremental fine-tuning (29 juin)", level=1)
    doc.add_paragraph(
        "Premier test de bout en bout du DAG blood_cell_incremental_finetune_pipeline "
        "depuis l'ajout de la planification et de la sélection automatique. PC Windows "
        "(GPU) rallumé pour l'occasion ; connexion SSH Airflow → Windows confirmée "
        "(authentification par clé OK)."
    )

    add_heading(doc, "3.1 Déroulé des tâches", level=2)
    add_table(
        doc,
        ["Tâche", "Résultat", "Durée"],
        [
            ["determine_batch_and_generation", "Succès — batch_001 / génération v94 "
             "calculés automatiquement", "< 1 s"],
            ["transfer_batch", "Succès — 200 images transférées vers le PC Windows (SFTP)", "14 s"],
            ["finetune_model", "Succès — fine-tuning + évaluation, nouvelle version "
             "MLflow v104 créée", "2,0 min"],
            ["check_promotion", "Succès — détecte que v104 ne dépasse pas la production", "< 1 s"],
            ["no_promotion", "Succès — branche correcte prise", "< 1 s"],
            ["promote_success / sync_to_datalake", "Skipped (attendu, branche non "
             "déclenchée)", "—"],
            ["pipeline_done", "Succès", "< 1 s"],
        ],
    )

    add_heading(doc, "3.2 Résultat du garde-fou de promotion", level=2)
    add_bullet(
        doc,
        " (v104, génération v94) : macro_f1 = 0,9978",
        bold_prefix="Nouveau challenger",
    )
    add_bullet(
        doc,
        " (v100, génération v90) : macro_f1 = 0,9983",
        bold_prefix="Production actuelle, inchangée",
    )
    add_bullet(
        doc,
        " — le garde-fou a fonctionné comme prévu : v104 reste "
        "@challenger, @production n'a pas été touché.",
        bold_prefix="0,9978 < 0,9983 → pas de promotion",
    )

    add_heading(doc, "3.3 Point d'attention", level=2)
    doc.add_paragraph(
        "Les 91 lots TIFF (data/tiff_batches/) avaient déjà tous été consommés "
        "manuellement par le passé (générations historiques v10 à v93, run_ids "
        "manual_batch_NNN_vMM_*). Le nouveau compteur automatique (Airflow Variable "
        "tiff_next_batch_idx) est reparti de zéro et a donc retraité batch_001 — pas "
        "de nouvelles données réelles, cohérent avec le caractère simulé du projet "
        "(aucune vraie donnée patient n'arrive). Le compteur est maintenant à 2 : la "
        "prochaine exécution planifiée (dimanche 4h) traitera batch_002."
    )

    add_heading(doc, "3.4 Conclusion", level=2)
    doc.add_paragraph(
        "Le mécanisme de planification + sélection automatique du lot/génération est "
        "validé en conditions réelles, de bout en bout, sans intervention manuelle "
        "(hormis le déclenchement de test et le rallumage du PC Windows). Aucune "
        "erreur rencontrée."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
