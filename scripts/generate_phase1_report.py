"""Génère le rapport Word Phase 1 dans reports/."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = Path(__file__).parents[1] / "reports" / "Phase1_etat_des_lieux.docx"

ROWS = [
    ("1", "Définir les objectifs & roadmap", "✅ Fait",
     "—",
     "Roadmap validée avec le mentor, 4 phases jusqu'au 7 juillet 2026"),
    ("2", "Environnement de développement reproductible", "✅ Fait",
     "Dockerfile, docker-compose.dev.yml,\n.env.example, requirements.txt",
     "Docker multi-services, .env pour toutes les variables sensibles"),
    ("3", "Collecter et pré-traiter les données", "✅ Fait",
     "src/data/dagshub_loader.py",
     "Téléchargement automatique depuis DagsHub, vérification par hash MD5"),
    ("4", "Créer une base de données (SQL)", "✅ Fait",
     "scripts/test_connections.py",
     "Supabase (PostgreSQL) — table predictions opérationnelle"),
    ("5", "Stocker les données via un script Python (one-shot)", "✅ Fait",
     "src/data/dagshub_loader.py",
     "Script exécuté une seule fois pour initialiser les données locales"),
    ("6", "Construire et évaluer un modèle ML de base", "✅ Fait",
     "src/train/training.py",
     "DenseNet-121, 8 classes, accuracy val + test, early stopping"),
    ("7", "Script training.py", "✅ Fait",
     "src/train/training.py",
     "Portable, piloté par .env, sauvegarde modèle + metrics.json"),
    ("8", "Script predict.py", "✅ Fait",
     "src/models/predict_model.py",
     "Inférence image → classe + confiance + top 3, flag classes critiques"),
    ("9", "API d'inférence — endpoint POST /training", "✅ Fait",
     "src/serving/api.py",
     "Lance l'entraînement, retourne val_acc + test_acc"),
    ("10", "API d'inférence — endpoint POST /predict", "✅ Fait",
     "src/serving/api.py",
     "Upload image → classe prédite + toutes les probabilités"),
]

HEADERS = ["#", "Livrable attendu", "Statut", "Fichier(s)", "Détail"]
COL_WIDTHS = [Cm(0.8), Cm(4.5), Cm(2.0), Cm(4.5), Cm(6.5)]


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


doc = Document()

# ── Styles globaux ────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Titre ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("Projet MLOps — Classification de Cellules Sanguines", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_heading("État des lieux — Phase 1 : Fondations", level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    f"Date : {datetime.date.today().strftime('%d %B %Y')}     "
    f"Deadline Phase 1 : 17 Juin 2026     "
    f"Branche Git : Sara"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── Tableau ───────────────────────────────────────────────────────────────────
table = doc.add_table(rows=1, cols=len(HEADERS))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

# En-tête
hdr_cells = table.rows[0].cells
for i, (cell, width, header) in enumerate(zip(hdr_cells, COL_WIDTHS, HEADERS)):
    cell.width = width
    set_cell_bg(cell, "1F497D")
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

# Lignes de données
for i, row_data in enumerate(ROWS):
    row_cells = table.add_row().cells
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    for j, (cell, content) in enumerate(zip(row_cells, row_data)):
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(content)
        run.font.size = Pt(9)
        if j == 2:  # colonne Statut
            run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
            run.bold = True

# ── Note de bas de page ───────────────────────────────────────────────────────
doc.add_paragraph()
note = doc.add_paragraph(
    "Note : La Phase 2 (Microservices, Suivi & Versioning — deadline 19 juin) "
    "intégrera le logging MLflow complet, le MLflow Registry et la comparaison "
    "automatique des modèles."
)
note.runs[0].font.size = Pt(9)
note.runs[0].font.italic = True
note.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.save(OUTPUT)
print(f"Rapport généré : {OUTPUT}")
