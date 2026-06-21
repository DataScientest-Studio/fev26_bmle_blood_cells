"""Génère reports/Fred/recap_airflow_mlflow.docx — récap technique du
pipeline d'entraînement distant (Airflow + SSH/Tailscale + MLflow).

Document de suivi de projet, pas committé sur GitHub (cf. convention des
autres récaps dans reports/) — à partager avec l'équipe par un canal interne
si besoin.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parents[1]
OUT = ROOT / "reports" / "Fred" / "recap_airflow_mlflow.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(18)
    return p


def main():
    doc = Document()

    title = doc.add_heading("Récap — Entraînement distant Airflow + SSH/Tailscale + MLflow", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Projet : fev26_bmle_blood_cells — {date.today():%d/%m/%Y}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 1. Contexte
    add_heading(doc, "1. Contexte et objectif", level=1)
    doc.add_paragraph(
        "Le Mac de Fred héberge Airflow et MLflow (Docker) mais n'a pas de GPU adapté à "
        "l'entraînement DenseNet-121. Le PC Windows de Fred, lui, a une RTX 4090. L'objectif "
        "était de faire piloter l'entraînement par Airflow (sur le Mac) tout en l'exécutant "
        "réellement sur le PC Windows, et de centraliser tous les résultats dans le même "
        "registry MLflow."
    )

    # 2. Architecture
    add_heading(doc, "2. Architecture mise en place", level=1)

    add_heading(doc, "2.1 Tailscale — réseau privé entre les deux machines", level=2)
    doc.add_paragraph(
        "Tailscale relie le Mac et le PC Windows par un réseau privé (VPN maillé), sans "
        "passer par Internet public — chaque machine a une IP fixe sur ce réseau privé "
        "(détails dans .env, non commités). Authentification du Mac vers le PC Windows par "
        "clé SSH dédiée (~/.ssh/airflow_to_windows), sans mot de passe."
    )

    add_heading(doc, "2.2 Airflow — orchestration", level=2)
    add_bullet(doc, "Provider apache-airflow-providers-ssh installé, connexion ssh_windows_gpu "
               "déclarée (host, user, clé privée — tout paramétré via variables d'environnement, "
               "rien en dur dans le code suivi par git, repo GitHub public).")
    add_bullet(doc, "DAG blood_cell_training_pipeline : la tâche train_model se connecte en SSH "
               "au PC Windows et y lance l'entraînement (venv Windows, GPU CUDA).")
    add_bullet(doc, "Le script distant (src/train/dl_crossval_train.py) enregistre lui-même le "
               "meilleur modèle dans le MLflow Registry et décide de la promotion — Airflow se "
               "contente ensuite de vérifier le résultat (tâche check_promotion).")
    add_bullet(doc, "Planification : chaque dimanche à 2h, ou déclenchement manuel à tout moment "
               "depuis l'interface.")

    add_heading(doc, "2.3 MLflow — registry et garde-fou de promotion", level=2)
    doc.add_paragraph(
        "Chaque entraînement (5 folds, cross-validation stratifiée) enregistre le meilleur fold "
        "dans le Model Registry, sous blood-cell-densenet121, avec un tag generation. La version "
        "n'est promue alias @production que si :"
    )
    add_bullet(doc, "le macro F1 ne régresse pas par rapport à la version @production actuelle,")
    add_bullet(doc, "ET aucun recall par classe ne régresse de plus de 2 points — sur les 8 "
               "classes (basophil, eosinophil, erythroblast, ig, lymphocyte, monocyte, "
               "neutrophil, platelet), pas seulement les 2 classes cliniquement critiques "
               "comme dans la version précédente du garde-fou.")
    doc.add_paragraph(
        "Si la nouvelle version ne passe pas ce garde-fou, elle reste @challenger et "
        "@production n'est pas touché — aucun risque de dégrader le modèle en prod avec un "
        "entraînement raté ou sur des données partielles."
    )

    add_heading(doc, "2.4 Convention de versionnage : generation", level=2)
    add_bullet(doc, "V0 : les 5 modèles DenseNet-121 (5-fold) déjà entraînés précédemment "
               "(stockés sur OneDrive), réenregistrés dans le Registry comme référence de "
               "départ — tag generation=v0, aucun alias touché.")
    add_bullet(doc, "V1, V2... : chaque nouveau cycle d'entraînement (nouvelles données ou "
               "réentraînement périodique) devient la génération suivante. C'est un tag "
               "logique, distinct du numéro de version MLflow auto-incrémenté.")

    # 3. État actuel
    add_heading(doc, "3. État actuel du Registry", level=1)
    add_bullet(doc, "V0 enregistrée : versions 1 à 5, tag generation=v0 (référence, 5 folds "
               "déjà entraînés précédemment).")
    add_bullet(doc, "V1 — test de validation (2 folds, 6 epochs) effectué avec succès cet "
               "après-midi : promu @production pour valider que toute la chaîne fonctionne "
               "(SSH → entraînement → registry → promotion).")
    add_bullet(doc, "Run complet V1 (5 folds, 20 epochs, ~1h30-2h sur la RTX 4090) lancé le "
               "soir même depuis Airflow — remplacera automatiquement la version @production "
               "si ses résultats sont au moins aussi bons (garde-fou décrit en 2.3).")

    # 4. Bugs corrigés
    add_heading(doc, "4. Problèmes préexistants corrigés au passage", level=1)
    doc.add_paragraph(
        "Plusieurs bugs ont été découverts en testant cette chaîne pour la première fois — "
        "aucun n'est lié au travail d'aujourd'hui, ils bloquaient déjà silencieusement "
        "certaines fonctionnalités :"
    )
    add_bullet(doc, "Le réseau Docker n'était pas nommé correctement (COMPOSE_PROJECT_NAME "
               "absent) — empêchait Airflow de démarrer si on suivait les commandes \"normales\".")
    add_bullet(doc, "Le DAG importait mlflow sans que ce package soit installé dans l'image "
               "Airflow — le DAG aurait été cassé dès le premier chargement.")
    add_bullet(doc, "Le Dockerfile MLflow utilisait --default-artifact-root au lieu de "
               "--artifacts-destination : ça empêchait register_model() de fonctionner pour "
               "tout client hors du conteneur (y compris depuis le Mac lui-même) — aucune "
               "promotion de modèle n'aurait jamais pu fonctionner.")
    add_bullet(doc, "models/.gitkeep était suivi par git en parallèle de models.dvc : DVC "
               "refusait de retracker le dossier tant que ce conflit existait, bloquant aussi "
               "silencieusement toute mise à jour du modèle vers le datalake.")

    # 5. Automatisation datalake + inférence
    add_heading(doc, "5. Automatisation : datalake DagsHub et conteneurs d'inférence", level=1)
    doc.add_paragraph(
        "Nouvelle tâche sync_to_datalake, exécutée automatiquement après chaque promotion "
        "@production confirmée :"
    )
    add_bullet(doc, "Téléchargement du modèle @production courant depuis le MLflow Registry.")
    add_bullet(doc, "dvc add + dvc push vers DagsHub (le datalake partagé de l'équipe) — "
               "models/best_DenseNet_121.pth contient maintenant les poids réellement servis "
               "en production, plus l'ancien modèle du 12 juin.")
    add_bullet(doc, "git commit + push automatique de models.dvc sur GitHub (token dédié, "
               "scope repo uniquement — voir section sécurité).")
    add_bullet(doc, "Redémarrage automatique de blood_cell_api / blood_cell_streamlit s'ils "
               "tournent, pour qu'ils rechargent le nouveau modèle (le chargement ne se fait "
               "qu'au démarrage du conteneur, jamais à chaud).")
    doc.add_paragraph(
        "Nécessite : un Dockerfile dédié pour Airflow (git, dvc, torch CPU, timm) et le socket "
        "Docker du Mac monté dans le conteneur Airflow (pour pouvoir redémarrer les conteneurs "
        "d'inférence)."
    )

    # 6. Sécurité
    add_heading(doc, "6. Sécurité — accès partagé avec Romane et Sara", level=1)
    doc.add_paragraph(
        "Avant cette mise en place, Airflow/MLflow n'étaient accessibles que depuis le Mac "
        "lui-même. Pour donner accès à Romane et Sara sans exposer le PC Windows ni donner un "
        "accès terminal/SSH à qui que ce soit, la recommandation retenue est le \"Share device\" "
        "de Tailscale (partage d'un appareil précis, pas une invitation au tailnet complet) : "
        "elles n'ont accès qu'au Mac, uniquement aux interfaces web Airflow et MLflow. Voir le "
        "guide de connexion séparé pour le détail."
    )
    add_bullet(doc, "Mot de passe admin/admin d'Airflow changé (n'était que le défaut de la "
               "doc officielle) — nouveau mot de passe donné dans le guide de connexion.")
    add_bullet(doc, "Token GitHub dédié au push automatique : classique, scope repo "
               "uniquement (pas admin:org) — minimise le risque en cas de fuite depuis le "
               "conteneur.")

    # 7. Suite
    add_heading(doc, "7. Prochaines étapes", level=1)
    add_bullet(doc, "Run complet V1 en cours ce soir — vérifier le lendemain que la version "
               "est bien promue @production et regarder les métriques par classe "
               "(recall_platelet en particulier).")
    add_bullet(doc, "Confirmer que le commit/push automatique de models.dvc s'est bien fait "
               "sur GitHub après la promotion.")
    add_bullet(doc, "À discuter en équipe : intégrer une source de données \"autre instrument\" "
               "(archive TCIA, format TIFF) pour les générations suivantes — limitation connue : "
               "cette source ne couvre que 7 classes sur 8 (pas de plaquettes).")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
