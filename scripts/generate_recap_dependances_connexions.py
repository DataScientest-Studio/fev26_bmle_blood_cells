"""Génère reports/Fred/recap_dependances_connexions.docx — inventaire des
dépendances externes (Supabase, DagsHub, GitHub, Tailscale, MLflow...) et des
connexions réseau nécessaires entre les différents services du projet.

Document de suivi de projet, pas committé sur GitHub (cf. convention des
autres récaps dans reports/) — à partager avec l'équipe par un canal interne
si besoin.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parents[1]
OUT = ROOT / "reports" / "Fred" / "recap_dependances_connexions.docx"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(18)
    return p


def add_table(doc, headers, rows, header_color="1F497D"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, hdr in zip(table.rows[0].cells, headers):
        set_cell_bg(cell, header_color)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)

    for i, values in enumerate(rows):
        row = table.add_row().cells
        bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        for cell, content in zip(row, values):
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            r = cell.paragraphs[0].add_run(content)
            r.font.size = Pt(8.5)
    return table


def main():
    doc = Document()

    title = doc.add_heading("Récap — Dépendances externes et connexions réseau", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Projet : fev26_bmle_blood_cells — {date.today():%d/%m/%Y}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 1. Services applicatifs
    add_heading(doc, "1. Services applicatifs (stack docker-compose.dev.yml)", level=1)
    add_table(
        doc,
        ["Service", "Rôle", "Port", "Dépend de"],
        [
            ["streamlit", "Interface utilisateur", "8501", "api, mlflow"],
            ["api (FastAPI)", "Backend ML (/predict, /feedback, /training)", "8000", "mlflow"],
            ["mlflow", "Tracking + Model Registry (SQLite)", "5001 -> 5000", "—"],
            ["preprocessing", "Téléchargement/préparation données (one-shot)", "—", "DagsHub"],
            ["training", "Entraînement DenseNet-121 one-shot", "—", "mlflow"],
        ],
    )
    doc.add_paragraph(
        "Connexion interne : streamlit -> api via http://api:8000, tous deux -> mlflow via "
        "http://mlflow:5000 (nom de service Docker). Vu depuis l'hôte/Tailscale : "
        "http://<IP>:5001."
    )

    # 2. Auth inter-services
    add_heading(doc, "2. Authentification inter-services", level=1)
    add_bullet(
        doc,
        " : clé partagée Streamlit <-> FastAPI, header X-API-Key (src/serving/api.py). "
        "Protège /predict, /feedback, /training. Si absente : auth désactivée (dev local/CI).",
        bold_prefix="API_SECRET_KEY",
    )
    add_bullet(
        doc,
        " : comptes utilisateurs stockés dans Supabase (src/auth/db.py, users.py) — pas de "
        "système d'auth séparé.",
        bold_prefix="Comptes utilisateurs",
    )

    # 3. Supabase
    add_heading(doc, "3. Supabase (PostgreSQL)", level=1)
    doc.add_paragraph(
        "Host : aws-0-eu-west-1.pooler.supabase.com:5432, pooler IPv4 (fonctionne en Docker, "
        "Windows, réseaux sans IPv6). 5 variables requises :"
    )
    add_code(doc, "SUPABASE_HOST / SUPABASE_PORT / SUPABASE_DB / SUPABASE_USER / SUPABASE_PASSWORD")
    doc.add_paragraph("Tables utilisées :")
    add_table(
        doc,
        ["Table", "Rôle"],
        [
            ["predictions", "Chaque appel /predict (scripts/init_db.py)"],
            ["prediction_feedback", "Désaccord médecin sur une prédiction"],
            ["dataset_images", "Métadonnées du dataset local"],
            ["training_runs", "Un run d'entraînement (src/monitoring/supabase_logger.py)"],
            ["class_metrics", "Métriques par classe d'un run"],
            ["confusion_matrices", "Matrice de confusion d'un run"],
            ["(table comptes)", "Authentification utilisateurs (src/auth/db.py)"],
        ],
        header_color="2E75B6",
    )
    doc.add_paragraph(
        "Utilisé par : api, streamlit, training, incremental_finetune.py, dl_crossval_train.py, "
        "et Airflow — qui transmet ces identifiants en variables d'environnement au process "
        "distant Windows (Airflow n'appelle pas load_dotenv(), donc supabase_env_exports() dans "
        "airflow/dags/_common.py construit les exports PowerShell correspondants)."
    )

    # 4. DagsHub
    add_heading(doc, "4. DagsHub — datalake (images + modèles)", level=1)
    doc.add_paragraph("Repo : Dumegan/Bloodcells-project. Variables :")
    add_code(doc, "DAGSHUB_USER / DAGSHUB_TOKEN (Settings > Tokens) / DAGSHUB_REPO_OWNER / DAGSHUB_REPO")
    doc.add_paragraph(
        "C'est le remote DVC : dvc push/dvc pull y stockent les vraies données (images, "
        "models/) ; Git ne garde que les fichiers .dvc (hash + métadonnées). Vérifié par "
        "scripts/check_connections.py."
    )

    # 5. GitHub
    add_heading(doc, "5. GitHub — push automatique post-promotion", level=1)
    doc.add_paragraph(
        "Token fine-grained, scope \"Contents: Read and write\" sur ce repo uniquement. Utilisé "
        "uniquement par sync_to_datalake() (airflow/dags/_common.py) : après promotion d'un "
        "modèle @production, Airflow commit+push models.dvc sur main avec ce token, puis "
        "redémarre blood_cell_api / blood_cell_streamlit (via /var/run/docker.sock monté dans "
        "le conteneur Airflow) pour qu'ils chargent la nouvelle version."
    )

    # 6. Tailscale
    add_heading(doc, "6. Tailscale — topologie réseau distribuée", level=1)
    add_bullet(
        doc,
        " : héberge Airflow + MLflow (serveur unique, SQLite). Romane et Sara s'y connectent "
        "via partage Tailscale \"Share device\" (accès à cet appareil seul) -> "
        "http://<MAC_TAILSCALE_IP>:8080 (Airflow) et :5001 (MLflow).",
        bold_prefix="Mac de Fred",
    )
    add_bullet(
        doc,
        " (GPU) : exécute les entraînements réels, déclenché à distance par Airflow via SSH.",
        bold_prefix="PC Windows de Fred",
    )
    doc.add_paragraph("Variables :")
    add_code(
        doc,
        "MAC_TAILSCALE_IP / WINDOWS_TAILSCALE_IP / WINDOWS_SSH_USER / WINDOWS_REPO_DIR / "
        "AIRFLOW_ADMIN_PASSWORD",
    )
    doc.add_paragraph(
        "Connexion Airflow définie dans docker-compose-airflow.yml : "
        "AIRFLOW_CONN_SSH_WINDOWS_GPU (clé privée ~/.ssh/airflow_to_windows montée en lecture "
        "seule dans le conteneur)."
    )

    # 7. Airflow
    add_heading(doc, "7. Airflow — orchestration (docker-compose-airflow.yml)", level=1)
    add_bullet(
        doc,
        " : Postgres dédié (airflow-db, credentials airflow/airflow), local au réseau Docker — "
        "pas Supabase.",
        bold_prefix="DB interne",
    )
    add_bullet(
        doc,
        " fev26_bmle_blood_cells_ml_network avec la stack docker-compose.dev.yml (doit déjà "
        "exister).",
        bold_prefix="Réseau externe partagé",
    )
    add_bullet(
        doc,
        " : Supabase, DagsHub, GitHub, Tailscale (toutes transmises au conteneur Airflow).",
        bold_prefix="Variables reçues",
    )

    # 8. DVC
    add_heading(doc, "8. DVC — versioning dataset (local + remote DagsHub)", level=1)
    doc.add_paragraph(
        "Fichiers .dvc trackés dans Git : Source_100.dvc, Source_full.dvc, runs/run1-5.dvc. "
        "Aucun .dvc pour data/tiff_batches/ actuellement — le tag dvc_dataset_hash y reste "
        "\"unknown\" pour les runs de fine-tuning sur ce lot."
    )

    # 9. MLflow
    add_heading(doc, "9. MLflow — tracking + Model Registry", level=1)
    doc.add_paragraph(
        "Backend SQLite (mlflow.db dans un volume Docker), artefacts sur disque local. Pas de "
        "remote partagé hors le Mac de Fred — c'est la seule instance faisant autorité (un run "
        "créé ailleurs, ex. sur la machine locale de quelqu'un d'autre, n'y apparaîtra jamais)."
    )

    # 10. Schéma des flux
    add_heading(doc, "10. Schéma des flux", level=1)
    add_code(
        doc,
        "Romane/Sara --Tailscale--> Mac Fred (Airflow :8080, MLflow :5001)\n"
        "                                |  SSH (Tailscale)\n"
        "                                v\n"
        "                       PC Windows Fred (GPU) -- entrainement reel\n"
        "                                |\n"
        "                  promotion @production (garde-fous macro_f1/recall)\n"
        "                                |\n"
        "          +---------------------+----------------------+\n"
        "          v                     v                      v\n"
        "   DagsHub (DVC)        GitHub (models.dvc)      Supabase (logs)\n"
        "                                |\n"
        "          docker restart blood_cell_api / blood_cell_streamlit",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
