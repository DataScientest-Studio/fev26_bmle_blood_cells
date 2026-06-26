"""Génère reports/Fred/recap_pptx_soutenance.docx — récap des erreurs trouvées
et des corrections apportées au pptx de soutenance (BloodCells_MLOps_Soutenance.pptx)
avant la soutenance du 7 juillet.

Document de suivi de projet, pas committé sur GitHub (cf. convention des
autres récaps dans reports/) — à partager avec l'équipe par un canal interne
si besoin.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parents[1]
OUT = ROOT / "reports" / "Fred" / "recap_pptx_soutenance.docx"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows, header_color="1F497D"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, hdr in zip(table.rows[0].cells, headers):
        set_cell_bg(cell, header_color)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
    for i, values in enumerate(rows):
        row = table.add_row().cells
        bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        for cell, content in zip(row, values):
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            r = cell.paragraphs[0].add_run(content)
            r.font.size = Pt(8.5)
    return table


def main():
    doc = Document()

    title = doc.add_heading("Récap — Corrections du pptx de soutenance", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"BloodCells_MLOps_Soutenance.pptx — {date.today():%d/%m/%Y}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_heading(doc, "Contexte", level=1)
    doc.add_paragraph(
        "Relecture du pptx de soutenance avant le 7 juillet, en croisant les chiffres "
        "affichés avec l'état réel du repo (manifestes .dvc, dossiers data/, dvc status)."
    )

    add_heading(doc, "Corrections appliquées dans le fichier", level=1)
    add_table(
        doc,
        ["Diapo", "Avant", "Après"],
        [
            ["9", "17 093 fichiers · jeu complet", "17 092 fichiers · jeu complet"],
            [
                "9",
                "run3 · run4 · run5.dvc — lots TIFF · 800 fichiers chacun",
                "run1 → run5.dvc — lots JPEG · 800 fichiers chacun",
            ],
            ["11", "KPI vide : « ~ ms »", "« ~42 ms »"],
            [
                "17 (doublon)",
                "Diapo brouillon dupliquée de la démo, coquille « Neutrophil »",
                "Supprimée",
            ],
        ],
    )
    doc.add_paragraph(
        "Le run3/4/5.dvc était à tort présenté comme du TIFF : ce sont en réalité des "
        "JPEG copiés depuis data/raw (script make_acquisition_runs.py, simulation de "
        "5 jours d'acquisition), pas les vraies images TIFF Matek/TCIA. Le tableau "
        "omettait aussi run1 et run2.dvc."
    )

    add_heading(doc, "Corrections faites à la main (PowerPoint)", level=1)
    add_bullet(
        doc,
        " : ajout de la ligne « 18 200 images TIFF (Matek/TCIA) » sous "
        "« 17 092 images · Mendeley PBC ».",
        bold_prefix="Diapo 1",
    )
    add_bullet(doc, " : ajout du pied de page manquant.", bold_prefix="Diapo 7")

    add_heading(doc, "Points identifiés, non corrigés", level=1)
    add_bullet(
        doc,
        " (91 lots × 200 = 18 200 images Matek/TCIA) n'est versionné ni par DVC ni par "
        "Git — juste des symlinks locaux vers l'archive externe. À trancher en équipe : "
        "faut-il les intégrer au datalake DVC ?",
        bold_prefix="data/tiff_batches/",
    )
    add_bullet(
        doc,
        " : data/Source_full.dvc déclare 17 093 fichiers, mais data/raw en contient "
        "17 092 en local — dvc status signale « modified: data/raw ». Indépendant du "
        "pptx, à corriger dans le repo.",
        bold_prefix="Écart DVC ↔ disque",
    )
    add_bullet(
        doc,
        " (« Data workflow ») reste sans pied de page, mais c'est une diapositive "
        "masquée (show=\"0\") — invisible en diaporama, donc non prioritaire.",
        bold_prefix="Diapo 8",
    )

    add_heading(doc, "Fichier", level=1)
    doc.add_paragraph(
        "~/Downloads/BloodCells_MLOps_Soutenance.pptx (corrigé, 27 diapos). "
        "Original pré-corrections sauvegardé localement avant édition."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
