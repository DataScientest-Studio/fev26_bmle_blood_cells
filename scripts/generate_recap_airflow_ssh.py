"""Génère reports/Fred/recap_airflow_ssh_windows.docx — récap pipeline
d'entraînement distant (Airflow + SSH/Tailscale + MLflow) et guide de
connexion pour Romane et Sara.

Ce fichier n'est PAS committé sur GitHub (repo public) : il contient les IP
Tailscale réelles. À transmettre à Romane et Sara par un canal privé
(email, Slack, OneDrive...).
"""

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

ROOT = Path(__file__).parents[1]
load_dotenv(ROOT / ".env")
OUT = ROOT / "reports" / "Fred" / "recap_airflow_ssh_windows.docx"

if not os.getenv("MAC_TAILSCALE_IP") or not os.getenv("WINDOWS_TAILSCALE_IP"):
    raise EnvironmentError(
        "MAC_TAILSCALE_IP et WINDOWS_TAILSCALE_IP doivent être définis dans ton .env local."
    )
MAC_IP = os.environ["MAC_TAILSCALE_IP"]
WINDOWS_IP = os.environ["WINDOWS_TAILSCALE_IP"]


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(18)
    return p


def main():
    doc = Document()

    title = doc.add_heading("Pipeline d'entraînement distant — Airflow + SSH/Tailscale + MLflow", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Projet : fev26_bmle_blood_cells — {date.today():%d/%m/%Y}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Document interne — contient des adresses réseau privées, ne pas publier.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0x33, 0x33)

    # 1. Contexte
    add_heading(doc, "1. Contexte et objectif", level=1)
    doc.add_paragraph(
        "Le Mac de Fred héberge Airflow et MLflow (Docker) mais n'a pas de GPU adapté à "
        "l'entraînement DenseNet-121. Le PC Windows de Fred, lui, a une RTX 4090. L'objectif "
        "était de faire piloter l'entraînement par Airflow (sur le Mac) tout en l'exécutant "
        "réellement sur le PC Windows, et de centraliser tous les résultats dans le même "
        "registry MLflow."
    )

    # 2. Architecture
    add_heading(doc, "2. Architecture mise en place", level=1)

    add_heading(doc, "2.1 Tailscale — réseau privé entre les deux machines", level=2)
    doc.add_paragraph(
        "Tailscale relie le Mac et le PC Windows par un réseau privé (VPN maillé), sans "
        "passer par Internet public. Chaque machine a une IP fixe sur ce réseau :"
    )
    add_bullet(doc, f"Mac (Airflow, MLflow) : {MAC_IP}")
    add_bullet(doc, f"PC Windows (GPU, entraînement) : {WINDOWS_IP}")
    doc.add_paragraph(
        "Authentification du Mac vers le PC Windows par clé SSH dédiée "
        "(~/.ssh/airflow_to_windows), sans mot de passe."
    )

    add_heading(doc, "2.2 Airflow — orchestration", level=2)
    add_bullet(doc, "Provider apache-airflow-providers-ssh installé, connexion ssh_windows_gpu "
               "déclarée (host, user, clé privée).")
    add_bullet(doc, "DAG blood_cell_training_pipeline : la tâche train_model se connecte en SSH "
               "au PC Windows et y lance l'entraînement (venv Windows, GPU CUDA).")
    add_bullet(doc, "Le script distant (src/train/dl_crossval_train.py) enregistre lui-même le "
               "meilleur modèle dans le MLflow Registry et décide de la promotion — Airflow se "
               "contente ensuite de vérifier le résultat (tâche check_promotion).")
    add_bullet(doc, "Planification : chaque dimanche à 2h, ou déclenchement manuel à tout moment "
               "depuis l'interface.")

    add_heading(doc, "2.3 MLflow — registry et garde-fou de promotion", level=2)
    doc.add_paragraph(
        "Chaque entraînement (5 folds, cross-validation stratifiée) enregistre le meilleur fold "
        "dans le Model Registry, sous blood-cell-densenet121, avec un tag generation. La version "
        "n'est promue alias @production que si :"
    )
    add_bullet(doc, "le macro F1 ne régresse pas par rapport à la version @production actuelle,")
    add_bullet(doc, "ET aucun recall par classe ne régresse de plus de 2 points — sur les 8 "
               "classes (basophil, eosinophil, erythroblast, ig, lymphocyte, monocyte, "
               "neutrophil, platelet), pas seulement les 2 classes cliniquement critiques "
               "comme dans la version précédente du garde-fou.")
    doc.add_paragraph(
        "Si la nouvelle version ne passe pas ce garde-fou, elle reste @challenger et "
        "@production n'est pas touché — aucun risque de dégrader le modèle en prod avec un "
        "entraînement raté ou sur des données partielles."
    )

    add_heading(doc, "2.4 Convention de versionnage : generation", level=2)
    add_bullet(doc, "V0 : les 5 modèles DenseNet-121 (5-fold) déjà entraînés précédemment "
               "(stockés sur OneDrive), réenregistrés dans le Registry comme référence de "
               "départ — tag generation=v0, aucun alias touché.")
    add_bullet(doc, "V1, V2... : chaque nouveau cycle d'entraînement (nouvelles données ou "
               "réentraînement périodique) devient la génération suivante. C'est un tag "
               "logique, distinct du numéro de version MLflow auto-incrémenté.")

    # 3. État actuel
    add_heading(doc, "3. État actuel du Registry", level=1)
    add_bullet(doc, "V0 enregistrée : versions 1 à 5, tag generation=v0 (référence).")
    add_bullet(doc, "V1 — test de validation (2 folds, 6 epochs) effectué avec succès : "
               "promu @production pour valider que toute la chaîne fonctionne (SSH → "
               "entraînement → registry → promotion).")
    add_bullet(doc, "Le run complet pour V1 (5 folds, 20 epochs, ~1h30-2h sur la RTX 4090) est "
               "prévu cette nuit, déclenché manuellement depuis Airflow.")

    # 4. Bugs corrigés
    add_heading(doc, "4. Problèmes préexistants corrigés au passage", level=1)
    doc.add_paragraph(
        "Trois bugs ont été découverts en testant cette chaîne pour la première fois — "
        "aucun n'est lié au travail d'aujourd'hui, ils bloquaient déjà silencieusement "
        "certaines fonctionnalités :"
    )
    add_bullet(doc, "Le réseau Docker n'était pas nommé correctement (COMPOSE_PROJECT_NAME "
               "absent) — empêchait Airflow de démarrer si on suivait les commandes \"normales\".")
    add_bullet(doc, "Le DAG importait mlflow sans que ce package soit installé dans l'image "
               "Airflow — le DAG aurait été cassé dès le premier chargement.")
    add_bullet(doc, "Le Dockerfile MLflow utilisait --default-artifact-root au lieu de "
               "--artifacts-destination : ça empêchait register_model() de fonctionner pour "
               "tout client hors du conteneur (y compris depuis le Mac lui-même) — aucune "
               "promotion de modèle n'aurait jamais pu fonctionner.")

    # 5. Connexion pour Romane et Sara
    add_heading(doc, "5. Comment se connecter (Romane, Sara)", level=1)
    doc.add_paragraph(
        "Principe de sécurité : Romane et Sara ne doivent avoir accès qu'au Mac (Airflow + "
        "MLflow), jamais au PC Windows, et jamais à un accès terminal/SSH — seulement aux "
        "interfaces web. Voir détail à l'étape 1."
    )

    add_heading(doc, "Étape 1 — Recevoir l'accès au Mac (pas au PC Windows)", level=2)
    doc.add_paragraph(
        "Fred partage uniquement le Mac via la fonction \"Share device\" de Tailscale "
        "(console admin Tailscale → Machines → le Mac → Share → ton adresse email) — pas une "
        "invitation au tailnet complet. Avec ce partage ciblé, tu n'as accès qu'à cet appareil "
        "précis : tu ne vois pas le PC Windows, tu ne peux rien atteindre d'autre sur le réseau "
        "de Fred. Installer ensuite l'application Tailscale (tailscale.com/download) et "
        "accepter le partage reçu par email."
    )
    doc.add_paragraph(
        "Tu n'as besoin d'aucune clé SSH ni d'accès terminal à quoi que ce soit — uniquement "
        "les interfaces web ci-dessous, dans ton navigateur."
    )

    add_heading(doc, "Étape 2 — Accéder à l'interface Airflow", level=2)
    add_code(doc, f"http://{MAC_IP}:8080")
    doc.add_paragraph("Identifiants : admin / admin")
    add_bullet(doc, "Vue \"Grid\" ou \"Graph\" du DAG blood_cell_training_pipeline : statut de "
               "chaque exécution (réussie / échouée / en cours).")
    add_bullet(doc, "Cliquer sur une tâche puis \"Logs\" pour voir le détail (notamment "
               "train_model, qui affiche la progression de l'entraînement sur le PC Windows).")

    add_heading(doc, "Étape 3 — Accéder à l'interface MLflow", level=2)
    add_code(doc, f"http://{MAC_IP}:5001")
    doc.add_paragraph(
        "Onglet \"Models\" → blood-cell-densenet121 : liste des versions, tags generation/fold, "
        "alias @production / @challenger. Onglet \"Experiments\" → blood_cell_crossval_ameliorees "
        "pour le détail des courbes d'entraînement par fold."
    )

    # 6. Lancer un nouvel entraînement
    add_heading(doc, "6. Lancer un nouvel entraînement manuellement", level=1)
    add_bullet(doc, "Dans l'interface Airflow, ouvrir le DAG blood_cell_training_pipeline.")
    add_bullet(doc, "Cliquer sur le bouton ▶ (\"Trigger DAG\") en haut à droite.")
    add_bullet(doc, "L'entraînement se lance automatiquement sur le PC Windows (GPU) — il faut "
               "que le PC Windows soit allumé et connecté à Tailscale pour que ça fonctionne.")
    add_bullet(doc, "Suivre la progression via les logs de la tâche train_model ; le résultat "
               "final (promu ou pas) apparaît dans la tâche check_promotion.")
    doc.add_paragraph(
        "Important : le PC Windows étant aussi utilisé pour jouer, lancer un entraînement "
        "complet (~1h30-2h) en même temps qu'une session de jeu ralentit fortement les deux "
        "(le GPU ne peut pas vraiment faire les deux en parallèle)."
    )

    # 7. Suite
    add_heading(doc, "7. Prochaines étapes", level=1)
    add_bullet(doc, "Lancement du run complet (5 folds, 20 epochs) cette nuit.")
    add_bullet(doc, "Vérifier demain que la version est bien promue @production et regarder "
               "les métriques par classe (recall_platelet en particulier).")
    add_bullet(doc, "À discuter en équipe : intégrer une source de données \"autre instrument\" "
               "(archive TCIA, format TIFF) pour les générations suivantes — limitation connue : "
               "cette source ne couvre que 7 classes sur 8 (pas de plaquettes).")
    add_bullet(doc, "Recommandé avant de partager l'accès à Romane et Sara : changer le mot de "
               "passe admin/admin d'Airflow (actuellement le défaut de la doc officielle) — "
               "airflow users create / airflow users reset-password.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
