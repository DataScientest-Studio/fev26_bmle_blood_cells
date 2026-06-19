from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_heading_color(paragraph, color: RGBColor):
    for run in paragraph.runs:
        run.font.color.rgb = color

def add_colored_heading(doc, text, level, color: RGBColor):
    p = doc.add_heading(text, level=level)
    set_heading_color(p, color)
    return p

def add_badge(paragraph, text, bg_color="2E86AB"):
    """Inline colored badge via XML shading on a run."""
    run = paragraph.add_run(f" {text} ")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_color)
    rPr.append(shd)
    return run

def add_file_row(table, path, description):
    row = table.add_row()
    row.cells[0].text = path
    row.cells[1].text = description
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.name = "Courier New"

def style_table(table):
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for cell in hdr.cells:
        cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1B4F72")
        cell._tc.tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

BLUE   = RGBColor(0x1B, 0x4F, 0x72)
GREEN  = RGBColor(0x1E, 0x8B, 0x4C)
ORANGE = RGBColor(0xD4, 0x6A, 0x00)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("Blood Cell Classification — ML Project")
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = BLUE

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run("Rapport de livrables — Phases 1 & 2")
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.add_run(f"Date : {date.today().strftime('%d/%m/%Y')}").font.size = Pt(11)

doc.add_page_break()

# ── SECTION 0 : Contexte ──────────────────────────────────────────────────────
add_colored_heading(doc, "Contexte du projet", 1, BLUE)
ctx = doc.add_paragraph(
    "Ce projet de formation BMLE (Blood Cell Machine Learning Engineering) a pour objectif "
    "de construire un système complet de classification automatique des cellules sanguines "
    "à partir d'images microscopiques. Le dataset comprend 8 classes de cellules : "
    "basophil, eosinophil, erythroblast, IG, lymphocyte, monocyte, neutrophil et platelet.\n\n"
    "L'architecture cible est un pipeline MLOps complet : acquisition des données, "
    "entraînement avec suivi d'expériences, API d'inférence, orchestration et monitoring."
)
ctx.style.font.size = Pt(10)

doc.add_paragraph()

# ── SECTION 1 : Phase 1 ───────────────────────────────────────────────────────
add_colored_heading(doc, "Phase 1 — Fondations", 1, BLUE)
p = doc.add_paragraph()
p.add_run("Statut : ")
add_badge(p, "TERMINÉ", "1E8B4C")
p.add_run("  Deadline : 17/06/2026")

doc.add_paragraph()

# --- 1.1 Environnement
add_colored_heading(doc, "1.1  Environnement de développement reproductible", 2, GREEN)

t = doc.add_table(rows=1, cols=2)
t.columns[0].width = Cm(7)
t.columns[1].width = Cm(9)
hdr = t.rows[0].cells
hdr[0].text = "Fichier"
hdr[1].text = "Rôle"
style_table(t)

for row in [
    ("requirements.txt",            "Dépendances Python de production"),
    ("requirements-dev.txt",        "Dépendances de développement (tests, lint)"),
    ("requirements-api.txt",        "Dépendances spécifiques à l'API FastAPI"),
    ("requirements-streamlit.txt",  "Dépendances de l'interface Streamlit"),
    ("setup.cfg",                   "Configuration flake8 et pytest"),
    (".env.example",                "Template des variables d'environnement"),
    ("Makefile",                    "Commandes raccourcies (train, test, lint…)"),
]:
    add_file_row(t, row[0], row[1])

doc.add_paragraph()

# --- 1.2 Données
add_colored_heading(doc, "1.2  Données collectées et pré-traitées", 2, GREEN)

t2 = doc.add_table(rows=1, cols=2)
t2.columns[0].width = Cm(7)
t2.columns[1].width = Cm(9)
hdr2 = t2.rows[0].cells
hdr2[0].text = "Fichier / Dossier"
hdr2[1].text = "Rôle"
style_table(t2)

for row in [
    ("src/data/make_dataset.py",           "Script de préparation et split du dataset"),
    ("src/data/dagshub_loader.py",         "Chargement des données depuis DagsHub/DVC"),
    ("src/data/source_100_manifest.json",  "Manifeste du dataset Source_100 (8 classes)"),
    ("Source_100.dvc",                     "Pointeur DVC — dataset 100 images/classe"),
    ("Source_full.dvc",                    "Pointeur DVC — dataset complet"),
    ("data/Source_100/",                   "Images organisées par classe (8 dossiers)"),
]:
    add_file_row(t2, row[0], row[1])

doc.add_paragraph()

# --- 1.3 Base de données SQL
add_colored_heading(doc, "1.3  Base de données SQL", 2, GREEN)

t3 = doc.add_table(rows=1, cols=2)
t3.columns[0].width = Cm(7)
t3.columns[1].width = Cm(9)
t3.rows[0].cells[0].text = "Fichier"
t3.rows[0].cells[1].text = "Rôle"
style_table(t3)

for row in [
    ("scripts/init_db.py",         "Initialisation des tables Supabase (one-shot)"),
    ("scripts/test_connections.py","Vérification de la connexion à la base de données"),
]:
    add_file_row(t3, row[0], row[1])

doc.add_paragraph()

# --- 1.4 Modèle ML
add_colored_heading(doc, "1.4  Modèle ML de base", 2, GREEN)

t4 = doc.add_table(rows=1, cols=2)
t4.columns[0].width = Cm(7)
t4.columns[1].width = Cm(9)
t4.rows[0].cells[0].text = "Fichier"
t4.rows[0].cells[1].text = "Rôle"
style_table(t4)

for row in [
    ("src/train/training.py",         "Script d'entraînement principal (DenseNet121)"),
    ("src/train/train_simple.py",     "Version simplifiée pour tests rapides"),
    ("src/models/train_model.py",     "Classe modèle et boucle d'entraînement"),
    ("src/models/predict_model.py",   "Inférence sur une image unique"),
    ("src/features/build_features.py","Transformations et augmentation de données"),
    ("configs/densenet121.yaml",      "Hyperparamètres du modèle DenseNet121"),
    ("models/best_DenseNet_121.pth",  "Poids du meilleur modèle entraîné"),
]:
    add_file_row(t4, row[0], row[1])

doc.add_paragraph()

# --- 1.5 API
add_colored_heading(doc, "1.5  API d'inférence", 2, GREEN)

t5 = doc.add_table(rows=1, cols=2)
t5.columns[0].width = Cm(7)
t5.columns[1].width = Cm(9)
t5.rows[0].cells[0].text = "Fichier"
t5.rows[0].cells[1].text = "Rôle"
style_table(t5)

for row in [
    ("src/serving/api.py",            "API FastAPI — endpoints POST /training et POST /predict"),
    ("src/serving/app.py",            "Interface utilisateur Streamlit"),
    ("src/serving/batch_inference.py","Inférence en lot sur un dossier d'images"),
]:
    add_file_row(t5, row[0], row[1])

doc.add_paragraph()

# --- 1.6 Tests
add_colored_heading(doc, "1.6  Tests automatisés", 2, GREEN)

t6 = doc.add_table(rows=1, cols=2)
t6.columns[0].width = Cm(7)
t6.columns[1].width = Cm(9)
t6.rows[0].cells[0].text = "Fichier"
t6.rows[0].cells[1].text = "Rôle"
style_table(t6)

for row in [
    ("tests/test_dataset.py",           "Tests de chargement et intégrité du dataset"),
    ("tests/test_predict.py",           "Tests de l'inférence (skippé si modèle absent en CI)"),
    ("tests/test_transforms.py",        "Tests des transformations d'images"),
    ("tests/fixtures/",                 "Images de test (2 par classe, 16 au total)"),
    (".github/workflows/ci.yml",        "Pipeline CI GitHub Actions (lint + tests)"),
]:
    add_file_row(t6, row[0], row[1])

doc.add_page_break()

# ── SECTION 2 : Phase 2 ───────────────────────────────────────────────────────
add_colored_heading(doc, "Phase 2 — Microservices, Suivi & Versionning", 1, BLUE)
p2 = doc.add_paragraph()
p2.add_run("Statut : ")
add_badge(p2, "TERMINÉ", "1E8B4C")
p2.add_run("  Deadline : 19/06/2026")

doc.add_paragraph()

# --- 2.1 MLflow
add_colored_heading(doc, "2.1  MLflow — Suivi d'expériences", 2, ORANGE)

t7 = doc.add_table(rows=1, cols=2)
t7.columns[0].width = Cm(7)
t7.columns[1].width = Cm(9)
t7.rows[0].cells[0].text = "Fichier / Dossier"
t7.rows[0].cells[1].text = "Rôle"
style_table(t7)

for row in [
    ("src/train/training.py",           "Logging MLflow : métriques, paramètres, artefacts"),
    ("src/evaluation/eval_best_models.py","Évaluation et marquage automatique du meilleur modèle"),
    ("src/evaluation/eval_experiments.py","Comparaison des runs MLflow"),
    ("models/.model_version",           "Fichier de tracking de version du modèle local"),
    ("mlruns/",                         "Base MLflow locale — expériences, runs, registry"),
]:
    add_file_row(t7, row[0], row[1])

doc.add_paragraph()

# --- 2.2 Versionning
add_colored_heading(doc, "2.2  Versionning données & modèles", 2, ORANGE)

t8 = doc.add_table(rows=1, cols=2)
t8.columns[0].width = Cm(7)
t8.columns[1].width = Cm(9)
t8.rows[0].cells[0].text = "Fichier"
t8.rows[0].cells[1].text = "Rôle"
style_table(t8)

for row in [
    ("Models.dvc",      "Pointeur DVC pour versionner les modèles entraînés"),
    ("Source_100.dvc",  "Pointeur DVC du dataset (synchronisé sur DagsHub)"),
]:
    add_file_row(t8, row[0], row[1])

doc.add_paragraph()
doc.add_paragraph(
    "5 runs d'acquisition simulés (8 classes × 100 images) ont été versionnés sur DagsHub "
    "via DVC Remote.",
    style="List Bullet"
)
doc.add_paragraph(
    "Le meilleur modèle est automatiquement comparé au modèle précédent à la fin de chaque run "
    "d'entraînement et promu dans le MLflow Registry si ses performances sont supérieures.",
    style="List Bullet"
)

doc.add_paragraph()

# --- 2.3 Docker
add_colored_heading(doc, "2.3  Containerisation (Docker)", 2, ORANGE)

t9 = doc.add_table(rows=1, cols=2)
t9.columns[0].width = Cm(7)
t9.columns[1].width = Cm(9)
t9.rows[0].cells[0].text = "Fichier"
t9.rows[0].cells[1].text = "Rôle"
style_table(t9)

for row in [
    ("docker/api/Dockerfile",          "Image Docker du service API FastAPI"),
    ("docker/mlflow/Dockerfile",       "Image Docker du serveur MLflow"),
    ("docker/streamlit/Dockerfile",    "Image Docker de l'interface Streamlit"),
    ("docker/docker-compose.dev.yml",  "Orchestration multi-services en développement"),
]:
    add_file_row(t9, row[0], row[1])

doc.add_page_break()

# ── SECTION 3 : Vue d'ensemble ────────────────────────────────────────────────
add_colored_heading(doc, "Vue d'ensemble de l'architecture", 1, BLUE)

doc.add_paragraph(
    "Le projet suit une structure de package Python standard :"
)

structure_lines = [
    "fev26_bmle_blood_cells/",
    "├── src/",
    "│   ├── data/          # Acquisition et préparation des données",
    "│   ├── features/      # Transformations et augmentation",
    "│   ├── models/        # Définition et inférence du modèle",
    "│   ├── train/         # Scripts d'entraînement",
    "│   ├── evaluation/    # Évaluation et comparaison MLflow",
    "│   └── serving/       # API FastAPI + Streamlit",
    "├── tests/             # Tests unitaires + fixtures",
    "├── scripts/           # Scripts utilitaires one-shot",
    "├── docker/            # Dockerfiles et docker-compose",
    "├── airflow/           # DAGs d'orchestration (Phase 3)",
    "├── configs/           # Fichiers de configuration YAML",
    "└── .github/workflows/ # Pipelines CI/CD",
]

p_code = doc.add_paragraph()
p_code.style = "No Spacing"
run = p_code.add_run("\n".join(structure_lines))
run.font.name = "Courier New"
run.font.size = Pt(9)

doc.add_paragraph()

# ── SECTION 4 : Phase suivante ────────────────────────────────────────────────
add_colored_heading(doc, "Phase 3 — Orchestration Airflow (à venir)", 1, BLUE)
p3 = doc.add_paragraph()
p3.add_run("Statut : ")
add_badge(p3, "EN COURS", "D46A00")

doc.add_paragraph()
doc.add_paragraph(
    "Les fondations sont posées. Le dossier airflow/ contient déjà :"
)
doc.add_paragraph("airflow/dags/blood_cell_pipeline.py — DAG Airflow du pipeline complet", style="List Bullet")
doc.add_paragraph("airflow/docker-compose-airflow.yml — stack Airflow en Docker", style="List Bullet")

doc.add_paragraph()
doc.add_paragraph(
    "Prochaines étapes : finaliser le DAG, connecter les tâches training → évaluation → "
    "promotion MLflow, et mettre en place le monitoring des prédictions en production."
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = Path(__file__).parents[1] / "rapport_livrables.docx"
doc.save(out)
print(f"Rapport généré : {out}")
