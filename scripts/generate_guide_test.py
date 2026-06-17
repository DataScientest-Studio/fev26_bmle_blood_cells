"""Génère le guide de test Phase 1 + état des lieux Phase 2 en Word dans reports/."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = Path(__file__).parents[1] / "reports" / "Sara" / "Guide_Test_Phases1_2.docx"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_code_block(doc, code: str):
    """Ajoute un bloc de code avec fond gris."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def add_info_box(doc, text: str, color="1F497D"):
    p = doc.add_paragraph()
    run = p.add_run(f"ℹ️  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_warning_box(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)


doc = Document()

# ── Style global ──────────────────────────────────────────────────────────────
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── Titre ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("Projet MLOps — Classification de Cellules Sanguines", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_heading("Guide de Test — Phase 1 (Windows & Mac)", level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    f"Date : {datetime.date.today().strftime('%d %B %Y')}     Deadline Phase 1 : 17 Juin 2026"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── Prérequis ─────────────────────────────────────────────────────────────────
doc.add_heading("Prérequis", level=2)

doc.add_paragraph("Prérequis 1 — Se placer dans le projet et activer l'environnement virtuel").runs[0].bold = True
doc.add_heading("Mac", level=3)
add_code_block(doc,
    "cd ~/fev26_bmle_blood_cells\n"
    "source .venv/bin/activate"
)
doc.add_heading("Windows (PowerShell)", level=3)
add_code_block(doc,
    "cd \"$env:USERPROFILE\\fev26_bmle_blood_cells\"\n"
    ".venv\\Scripts\\Activate.ps1"
)

doc.add_paragraph("Prérequis 2 — Fichier .env configuré").runs[0].bold = True
p = doc.add_paragraph("Copier ")
p.add_run(".env.example").font.name = "Courier New"
p.add_run(" en ")
p.add_run(".env").font.name = "Courier New"
p.add_run(" et remplir les valeurs — en particulier ")
r = p.add_run("SUPABASE_PASSWORD")
r.font.name = "Courier New"
r.bold = True
p.add_run(" et ")
r2 = p.add_run("DAGSHUB_TOKEN")
r2.font.name = "Courier New"
r2.bold = True
p.add_run(".")

# ── ÉTAPE PRÉALABLE ───────────────────────────────────────────────────────────
doc.add_heading("Étape préalable — Initialisation Supabase (déjà fait ✅)", level=2)
add_info_box(doc, "Cette étape a déjà été réalisée le 16 juin 2026. "
             "Les tables 'predictions' et 'dataset_images' existent dans Supabase. "
             "Ne pas relancer sauf en cas de recréation complète de la base.")
doc.add_paragraph("Pour mémoire, la commande était :")
doc.add_heading("Terminal Mac", level=3)
add_code_block(doc,
    "cd ~/fev26_bmle_blood_cells\n"
    "source .venv/bin/activate\n"
    "python scripts/init_db.py"
)
doc.add_heading("Terminal Windows (PowerShell)", level=3)
add_code_block(doc,
    "cd \"$env:USERPROFILE\\fev26_bmle_blood_cells\"\n"
    ".venv\\Scripts\\Activate.ps1\n"
    "python scripts/init_db.py"
)
doc.add_paragraph("Résultat attendu :").runs[0].bold = True
add_code_block(doc,
    "Connexion à Supabase...\n"
    "  [OK] Connexion établie\n\n"
    "Création des tables...\n"
    "  [OK] Tables 'predictions' et 'dataset_images' créées (ou déjà existantes)\n\n"
    "Terminé."
)

# ── TEST 1 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 1 — Connexions infrastructure (Supabase + DagsHub)", level=2)
add_info_box(doc, "Ce test vérifie les deux connexions infrastructure en une seule commande.")
doc.add_heading("Terminal Mac", level=3)
add_code_block(doc,
    "python scripts/test_connections.py"
)
doc.add_heading("Terminal Windows (PowerShell)", level=3)
add_code_block(doc,
    "python scripts/test_connections.py"
)
doc.add_paragraph("Résultat attendu :").runs[0].bold = True
add_code_block(doc,
    "==================================================\n"
    "  Test des connexions infrastructure\n"
    "==================================================\n\n"
    "[1/2] Supabase (PostgreSQL)\n"
    "  [OK] Connexion établie\n"
    "  [OK] Insertion OK (id=...)\n"
    "  [OK] Lecture OK → ('test_connection.jpg', 'neutrophil', 0.99)\n"
    "  [OK] Nettoyage de la ligne de test OK\n\n"
    "[2/2] DagsHub (datalake images + modèles)\n"
    "  [OK] Manifest Models.dvc OK (md5=xxxxxxxx...)\n"
    "  [OK] Manifest Source_100.dvc OK (md5=xxxxxxxx...)\n\n"
    "==================================================\n"
    "  Supabase     OK\n"
    "  DagsHub      OK\n"
    "=================================================="
)

# ── TEST 2 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 2 — Connexion DagsHub (téléchargement)", level=2)
doc.add_heading("VSCode (Windows & Mac)", level=3)
for step in [
    "Ouvrir Run & Debug (Ctrl+Shift+D / Cmd+Shift+D)",
    'Sélectionner "Test DagsHub connection"',
    "Cliquer sur ▶️",
]:
    doc.add_paragraph(step, style="List Number")

doc.add_paragraph("Résultat attendu :").runs[0].bold = True
add_code_block(doc,
    "[1/2] Modèle DenseNet\n"
    "  Statut  : déjà à jour\n"
    "  Présent : True\n"
    "  Taille  : 27.2 MB\n\n"
    "[2/2] Source_100 (images de test)\n"
    "  Statut  : déjà à jour\n"
    "  Images  : 100/100\n\n"
    "Connexion DagsHub OK"
)

# ── TEST 3 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 3 — Entraînement (training.py sans API)", level=2)
doc.add_heading("VSCode (Windows & Mac)", level=3)
for step in [
    'Run & Debug → sélectionner "Test training (Source_100 - 1 epoch)"',
    "Cliquer sur ▶️ — durée ~1-2 minutes",
]:
    doc.add_paragraph(step, style="List Number")

doc.add_paragraph("Résultat attendu :").runs[0].bold = True
add_code_block(doc,
    "Device  : cpu (ou mps sur Mac M1/M2)\n"
    "Images  : 100 dans 8 classes\n"
    "Split   : train=69  val=16  test=15\n\n"
    "Phase 1 — backbone gelé (1 epoch)\n"
    "  Ep 01  train=0.xxx  val=0.xxx  (Xs)\n\n"
    "Phase 2 — fine-tuning (1 epoch, patience=3)\n"
    "  Ep 01  train=0.xxx  val=0.xxx  (Xs)\n\n"
    "Meilleur val_acc : 0.xxxx\n"
    "Test accuracy    : 0.xxxx\n"
    "Modèle           : models/best_densenet121.pth\n"
    "MLflow run ID    : xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx\n\n"
    "MLflow Registry :\n"
    "  Premier modèle enregistré → promu en Production\n\n"
    "Résumé : val_acc=0.xxxx  test_acc=0.xxxx"
)
add_info_box(doc, "Avec 100 images et 1 epoch, les métriques seront faibles (~0.10-0.25). "
             "Ce test vérifie que le pipeline fonctionne, pas la qualité du modèle.")

doc.add_heading("Terminal Mac", level=3)
add_code_block(doc,
    "cd ~/fev26_bmle_blood_cells\n"
    "source .venv/bin/activate\n"
    "python -m src.train.training \\\n"
    "  --data-dir data/Source_100 --output-dir models \\\n"
    "  --epochs-head 1 --epochs-full 1 --batch-size 8"
)
doc.add_heading("Terminal Windows (PowerShell)", level=3)
add_code_block(doc,
    "cd \"$env:USERPROFILE\\fev26_bmle_blood_cells\"\n"
    ".venv\\Scripts\\Activate.ps1\n"
    "python -m src.train.training `\n"
    "  --data-dir data/Source_100 --output-dir models `\n"
    "  --epochs-head 1 --epochs-full 1 --batch-size 8"
)

# ── TEST 4 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 4 — Prédiction (predict_model.py sans API)", level=2)
add_warning_box(doc, "Nécessite que le Test 2 ait été lancé une fois (génère models/best_densenet121.pth).")
add_info_box(doc, "Avec le modèle test rapide (1 epoch), la prédiction peut être incorrecte — c'est attendu.")

doc.add_heading("VSCode (Windows & Mac)", level=3)
for step in [
    'Run & Debug → sélectionner "Test predict_model (une image)"',
    "Cliquer sur ▶️",
]:
    doc.add_paragraph(step, style="List Number")

doc.add_paragraph("Résultat attendu :").runs[0].bold = True
add_code_block(doc,
    "Prédiction : NEUTROPHIL  (ou autre classe si modèle test rapide)\n"
    "Confiance  : xx.x%\n\n"
    "Top 3 :\n"
    "  neutrophil      xx.x%\n"
    "  lymphocyte      xx.x%\n"
    "  monocyte        xx.x%"
)
doc.add_heading("Terminal Mac", level=3)
add_code_block(doc,
    "python -m src.models.predict_model \\\n"
    "  --image data/Source_100/neutrophil/BNE_100878.jpg \\\n"
    "  --model models/best_densenet121.pth"
)
doc.add_heading("Terminal Windows (PowerShell)", level=3)
add_code_block(doc,
    "python -m src.models.predict_model `\n"
    "  --image data/Source_100/neutrophil/BNE_100878.jpg `\n"
    "  --model models/best_densenet121.pth"
)

# ── TEST 5 ────────────────────────────────────────────────────────────────────
doc.add_heading("Test 5 — API FastAPI complète", level=2)
doc.add_heading("Étape 4.1 — Démarrer le serveur", level=3)
doc.add_paragraph("VSCode (Windows & Mac) :").runs[0].bold = True
for step in [
    'Run & Debug → sélectionner "Lancer API FastAPI"',
    "Cliquer sur ▶️",
    "Attendre : INFO: Application startup complete.",
]:
    doc.add_paragraph(step, style="List Number")

add_info_box(doc, "L'erreur 'size mismatch' au démarrage est corrigée. "
             "Le modèle est chargé automatiquement si models/best_densenet121.pth existe.")

doc.add_heading("Terminal Mac", level=3)
add_code_block(doc, "source .venv/bin/activate\nuvicorn src.serving.api:app --reload --port 8000")
doc.add_heading("Terminal Windows (PowerShell)", level=3)
add_code_block(doc, ".venv\\Scripts\\Activate.ps1\nuvicorn src.serving.api:app --reload --port 8000")

doc.add_heading("Étape 4.2 — Swagger UI (http://localhost:8000/docs)", level=3)
add_info_box(doc, "Les endpoints sont maintenant groupés en 3 sections dans Swagger.")
add_code_block(doc,
    "Inférence      →  POST /predict\n"
    "Entraînement   →  POST /training\n"
    "Info           →  GET /health  |  GET /  |  GET /classes  |  GET /model-info"
)

tests_swagger = [
    ("GET /health (groupe Info)", "Try it out → Execute",
     '{ "status": "ok" }'),
    ("POST /predict (groupe Inférence)", "Try it out → Choose File → image dans data/Source_100/ → Execute",
     '{ "class": "Neutrophil", "confidence": 0.876, "all_probas": {...} }'),
    ("POST /training (groupe Entraînement)",
     'Try it out → body : {"data_dir":"data/Source_100","epochs_head":1,"epochs_full":1,"batch_size":8} → Execute (~2 min)',
     '{ "status": "ok", "val_acc": 0.xxxx, "test_acc": 0.xxxx, "model_path": "..." }'),
]
for name, action, response in tests_swagger:
    doc.add_paragraph(name).runs[0].bold = True
    doc.add_paragraph(action)
    doc.add_paragraph("Réponse attendue :").runs[0].italic = True
    add_code_block(doc, response)

doc.add_heading("Étape 4.3 — Terminal (optionnel)", level=3)
doc.add_heading("Mac (curl)", level=4)
add_code_block(doc,
    "curl http://localhost:8000/health\n\n"
    "curl -X POST http://localhost:8000/predict \\\n"
    "  -F \"file=@data/Source_100/neutrophil/BNE_100878.jpg\"\n\n"
    "curl -X POST http://localhost:8000/training \\\n"
    "  -H \"Content-Type: application/json\" \\\n"
    "  -d '{\"data_dir\":\"data/Source_100\",\"epochs_head\":1,\"epochs_full\":1,\"batch_size\":8}'"
)
doc.add_heading("Windows (PowerShell)", level=4)
add_code_block(doc,
    "Invoke-RestMethod -Uri \"http://localhost:8000/health\" -Method GET\n\n"
    "$form = @{ file = Get-Item \"data\\Source_100\\neutrophil\\BNE_100878.jpg\" }\n"
    "Invoke-RestMethod -Uri \"http://localhost:8000/predict\" -Method POST -Form $form\n\n"
    "$body = @{data_dir=\"data/Source_100\";epochs_head=1;epochs_full=1;batch_size=8} | ConvertTo-Json\n"
    "Invoke-RestMethod -Uri \"http://localhost:8000/training\" -Method POST -Body $body -ContentType \"application/json\""
)

# ── TABLEAU RÉCAP ─────────────────────────────────────────────────────────────
doc.add_heading("Récapitulatif des tests Phase 1", level=2)
recap_rows = [
    ("init_db.py",      "python scripts/init_db.py",              "Tables 'predictions' et 'dataset_images' créées (ou déjà existantes)"),
    ("Connexions",      "python scripts/test_connections.py",      "Supabase OK + DagsHub OK"),
    ("DagsHub DL",      '"Test DagsHub connection" (VSCode)',      "Connexion DagsHub OK + modèle présent"),
    ("training.py",     '"Test training (Source_100 - 1 epoch)"', "Résumé : val_acc=... + MLflow run ID affichés"),
    ("predict_model.py",'"Test predict_model (une image)"',       "Classe + confiance affichés"),
    ("API démarrage",   '"Lancer API FastAPI"',                   "Application startup complete sans erreur"),
    ("API /health",     "Swagger → GET /health",                  '{"status":"ok"}'),
    ("API /predict",    "Swagger → POST /predict",                "class + confidence retournés"),
    ("API /training",   "Swagger → POST /training",               "val_acc + test_acc retournés"),
]
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdrs = ["Test", "Commande VSCode", "Succès si..."]
for i, (cell, hdr) in enumerate(zip(table.rows[0].cells, hdrs)):
    set_cell_bg(cell, "1F497D")
    set_cell_border(cell)
    run = cell.paragraphs[0].add_run(hdr)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

for i, row_data in enumerate(recap_rows):
    row = table.add_row().cells
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    for j, (cell, content) in enumerate(zip(row, row_data)):
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        run = cell.paragraphs[0].add_run(content)
        run.font.size = Pt(9)
        if j == 2:
            run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ÉTAT DES LIEUX — PHASE 2
# ══════════════════════════════════════════════════════════════════════════════
phase2_title = doc.add_heading("État des lieux — Phase 2 : MLflow Tracking", level=1)
phase2_title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph(
    f"Rédigé le {datetime.date.today().strftime('%d %B %Y')}  —  "
    "Deadline Phase 2 : 1er Juillet 2026"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── Objectifs Phase 2 ─────────────────────────────────────────────────────────
doc.add_heading("Objectifs atteints", level=2)
p = doc.add_paragraph()
for line in [
    "Tracking MLflow complet sur chaque run d'entraînement (métriques, paramètres, artifacts, tags)",
    "Model Registry avec aliases @production / @challenger (MLflow 3.x — stages dépréciés)",
    "Garde-fous de promotion pour les classes critiques erythroblast et ig",
    "predict_model.py câblé au Registry MLflow (@production) + logging Supabase",
    "MLproject pour la reproductibilité des runs",
    "Serveur MLflow dans Docker (port 5001, SQLite, artifacts servis via HTTP)",
]:
    p = doc.add_paragraph(line, style="List Bullet")
    p.runs[0].font.size = Pt(10)

add_info_box(doc, "Ni Airflow ni Cron ne sont utilisés en Phase 2. "
             "Le scheduling automatique des retrains est prévu en Phase 3.")

# ── Tableau État des lieux ────────────────────────────────────────────────────
doc.add_heading("Composants Phase 2 — État des lieux", level=2)

etat_rows = [
    # (Composant, Fichier, Statut, Description)
    ("Serveur MLflow",      "docker/mlflow/Dockerfile\ndocker/docker-compose.dev.yml",
     "OK", "Container Docker port 5001 — mlflow==3.11.1\n--serve-artifacts activé"),
    ("Tracking training",   "src/train/training.py",
     "OK", "10 params + 11 métriques + 3 artifacts + 2 tags\nloggés à chaque run"),
    ("Artifacts loggés",    "models/confusion_matrix.png\nmodels/classification_report.txt\nmodels/label_mapping.json",
     "OK", "Uploadés via HTTP vers le container\nVisibles dans MLflow UI → Artifacts"),
    ("Model Registry",      "MLFLOW_MODEL_NAME\n= blood-cell-densenet121",
     "OK", "Aliases @production / @challenger\n(MLflow 3.x — pas de stages)"),
    ("Garde-fous promotion","_register_and_promote()\nRECALL_TOLERANCE = 0.02",
     "OK", "macro_f1 >= prod + recall_erythroblast\net recall_ig non régressifs (±2%)"),
    ("Prédiction MLflow",   "src/models/predict_model.py",
     "OK", "Charge @production depuis Registry\nFallback .pth si MLflow indisponible"),
    ("Logging Supabase",    "predict_model._log_to_supabase()",
     "OK", "inference_ms + predicted_class\nInsert dans table predictions"),
    ("MLproject",           "MLproject + conda.yaml",
     "OK", "mlflow run . -e train --env-manager=local\n-P data_dir=... -P epochs_head=..."),
    ("Tables SQL MLflow",   "mlruns/mlflow.db (SQLite)",
     "AUTO", "Créées automatiquement par Alembic\nau démarrage du container"),
    ("API /predict MLflow", "src/serving/api.py",
     "Phase 3", "Non câblée — charge encore le .pth local\nRefactoring prévu Phase 3"),
    ("Scheduling retrain",  "Airflow / Cron",
     "Phase 3", "Aucun scheduler en Phase 2\nA définir avec le mentor"),
]

table2 = doc.add_table(rows=1, cols=4)
table2.style = "Table Grid"
hdrs2 = ["Composant", "Fichier / Config", "Statut", "Description"]
col_colors = ["1F497D", "1F497D", "1F497D", "1F497D"]
for i, (cell, hdr) in enumerate(zip(table2.rows[0].cells, hdrs2)):
    set_cell_bg(cell, "1F497D")
    set_cell_border(cell)
    r = cell.paragraphs[0].add_run(hdr)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)

STATUS_COLORS = {
    "OK":      ("D4EDDA", RGBColor(0x00, 0x70, 0x00)),
    "AUTO":    ("FFF3CD", RGBColor(0x85, 0x60, 0x00)),
    "Phase 3": ("F8D7DA", RGBColor(0x80, 0x00, 0x00)),
}

for i, (composant, fichier, statut, description) in enumerate(etat_rows):
    row = table2.add_row().cells
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    status_bg, status_color = STATUS_COLORS.get(statut, (bg, RGBColor(0, 0, 0)))

    for j, (cell, content) in enumerate(zip(row, [composant, fichier, statut, description])):
        cell_bg = status_bg if j == 2 else bg
        set_cell_bg(cell, cell_bg)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(content)
        r.font.size = Pt(8.5)
        if j == 2:
            r.bold = True
            r.font.color.rgb = status_color

# ── Tables SQL détail ─────────────────────────────────────────────────────────
doc.add_heading("Tables SQL MLflow (auto-créées par Alembic)", level=2)
add_info_box(doc, "Contrairement à Supabase (init_db.py manuel), les tables MLflow "
             "sont créées automatiquement par le framework au premier démarrage du container.")

sql_rows = [
    ("experiments",           "Expérience bloodcells-densenet121"),
    ("runs",                  "Chaque run training.py (run_id, status, timestamps)"),
    ("params",                "batch_size, lr_head, n_train... (mlflow.log_param)"),
    ("metrics",               "macro_f1, recall_ig... (mlflow.log_metric)"),
    ("latest_metrics",        "Dernière valeur de chaque métrique par run"),
    ("tags",                  "git_commit, run_type"),
    ("registered_models",     "blood-cell-densenet121"),
    ("model_versions",        "v1 → v8 (une par run)"),
    ("registered_model_aliases", "@production, @challenger"),
    ("logged_models",         "Modèles loggés via mlflow.pytorch.log_model()"),
    ("alembic_version",       "Version du schéma SQLite (migrations auto)"),
]

table3 = doc.add_table(rows=1, cols=2)
table3.style = "Table Grid"
for cell, hdr in zip(table3.rows[0].cells, ["Table SQLite", "Rôle dans le projet"]):
    set_cell_bg(cell, "2E75B6")
    set_cell_border(cell)
    r = cell.paragraphs[0].add_run(hdr)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9)

for i, (table_name, role) in enumerate(sql_rows):
    row = table3.add_row().cells
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    for j, (cell, content) in enumerate(zip(row, [table_name, role])):
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(content)
        r.font.size = Pt(9)
        if j == 0:
            r.font.name = "Courier New"

# ── Métriques loggées ─────────────────────────────────────────────────────────
doc.add_heading("Métriques et paramètres loggés dans chaque run", level=2)

doc.add_paragraph("Paramètres (10)").runs[0].bold = True
add_code_block(doc,
    "batch_size · lr_head · lr_full · weight_decay · epochs_head · epochs_full\n"
    "n_train · n_val · n_test · optimizer · part"
)
doc.add_paragraph("Métriques (11 + courbes par epoch)").runs[0].bold = True
add_code_block(doc,
    "best_val_acc · test_acc · macro_f1 · weighted_f1\n"
    "precision_macro · recall_macro\n"
    "recall_erythroblast  <- classe critique\n"
    "recall_ig            <- classe critique\n"
    "train_time_s · n_params\n"
    "train_loss_epN · val_loss_epN · train_acc_epN · val_acc_epN  (N = epoch)"
)
doc.add_paragraph("Tags (2)").runs[0].bold = True
add_code_block(doc, "git_commit  ·  run_type (base | retrain)")

# ── Garde-fous ────────────────────────────────────────────────────────────────
doc.add_heading("Logique des garde-fous de promotion", level=2)
add_code_block(doc,
    "Nouveau run -> @challenger (toujours)\n"
    "       |\n"
    "       v\n"
    "  macro_f1 >= prod_f1                      ?\n"
    "  recall_erythroblast >= prod_ery - 0.02   ?   <- tolérance 2%\n"
    "  recall_ig >= prod_ig - 0.02              ?   <- tolérance 2%\n"
    "       |\n"
    "   OUI (tout)        NON (un seul suffit)\n"
    "  @production   ->  reste @challenger + raison affichée"
)
add_info_box(doc,
    "Script de démonstration : python scripts/demo_garde_fou.py\n"
    "Simule un challenger dégradé et affiche la décision du garde-fou avec les valeurs réelles."
)

doc.save(OUTPUT)
print(f"Guide généré : {OUTPUT}")
